#!/usr/bin/env python3
"""
build-nbais-wc-template.py

Generates templates/verticals/nbais-wc/master.docx using python-docx.
This is the NBAIS Workers' Compensation proposal master template.

Usage: python3 scripts/build-nbais-wc-template.py
"""

# NOTE: This script generates master.docx locally.
# Template assets are loaded from S3 at runtime by the proposal-generator service.
# Run with --sync to push changes to S3 after generation.

import argparse
import os
import subprocess
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Paths ──────────────────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
REPO_ROOT = os.path.dirname(SCRIPT_DIR)
OUTDIR = os.path.join(REPO_ROOT, 'templates', 'verticals', 'nbais-wc')
OUTPATH = os.path.join(OUTDIR, 'master.docx')
LOGO_H_PATH = os.path.join(OUTDIR, 'logo_horizontal.png')
LOGO_S_PATH = os.path.join(OUTDIR, 'logo_stacked.png')

# ─── Colors ─────────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x1F, 0x38, 0x64)
BLUE      = RGBColor(0x2E, 0x75, 0xB6)
LT_BLUE   = RGBColor(0xEB, 0xF3, 0xFB)
GRAY      = RGBColor(0x59, 0x59, 0x59)
MID_GRAY  = RGBColor(0xB0, 0xB0, 0xB0)
LT_GRAY   = RGBColor(0xF5, 0xF5, 0xF5)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

NAVY_HEX     = '1F3864'
BLUE_HEX     = '2E75B6'
LT_BLUE_HEX  = 'EBF3FB'
GRAY_HEX     = '595959'
MID_GRAY_HEX = 'B0B0B0'
LT_GRAY_HEX  = 'F5F5F5'
BORDER_HEX   = 'CCCCCC'
WHITE_HEX    = 'FFFFFF'

# ─── Dimensions ─────────────────────────────────────────────────────────────
PAGE_W       = 12240   # 8.5in in twips
PAGE_H       = 15840   # 11in in twips
L_MARGIN     = 792     # 0.55in
R_MARGIN     = 792     # 0.55in
T_MARGIN     = 720     # 0.5in
B_MARGIN     = 720     # 0.5in
CONTENT_W    = 9360   # 6.5in at 1-inch margins (8.5in - 1in - 1in)
HDR_DIST     = 432     # 0.3in

FONT = 'Arial'

# ─────────────────────────────────────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────────────────────────────────────

def set_font(run, name=FONT, size_pt=10, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_para_shading(para, fill_hex):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def set_para_bottom_border(para, color_hex, size_eighths_pt):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size_eighths_pt))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_para_top_border(para, color_hex, size_eighths_pt):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), str(size_eighths_pt))
    top.set(qn('w:space'), '4')
    top.set(qn('w:color'), color_hex)
    pBdr.append(top)
    pPr.append(pBdr)


def add_right_tab_stop(para, pos_twips):
    pPr = para._p.get_or_add_pPr()
    tabs_el = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(pos_twips))
    tabs_el.append(tab)
    pPr.append(tabs_el)


def set_table_width(tbl, width_twips):
    tblPr = tbl._tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


def set_cell_width(cell, width_twips):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def set_cell_margins(cell, top=80, bottom=80, left=115, right=115):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def set_row_height(row, height_twips, exact=False):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trH = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(height_twips))
    if exact:
        trH.set(qn('w:hRule'), 'exact')
    trPr.append(trH)


def set_row_header(row):
    """Enable header row repeat on page break."""
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), '1')
    trPr.append(tblHeader)


def set_table_borders(tbl, color=BORDER_HEX, size=6, sides=None):
    if sides is None:
        sides = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
    tblPr = tbl._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(size))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)


def remove_table_borders(tbl):
    tblPr = tbl._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_cell_border(cell, sides):
    """sides: dict of side -> {val, sz, color}"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, spec in sides.items():
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), spec.get('val', 'single'))
        el.set(qn('w:sz'), str(spec.get('sz', 6)))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), spec.get('color', BORDER_HEX))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_no_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_table_alignment(tbl, alignment=WD_TABLE_ALIGNMENT.CENTER):
    tblPr = tbl._tbl.tblPr
    jc = OxmlElement('w:jc')
    if alignment == WD_TABLE_ALIGNMENT.CENTER:
        jc.set(qn('w:val'), 'center')
    elif alignment == WD_TABLE_ALIGNMENT.LEFT:
        jc.set(qn('w:val'), 'left')
    elif alignment == WD_TABLE_ALIGNMENT.RIGHT:
        jc.set(qn('w:val'), 'right')
    tblPr.append(jc)


def set_section_margins(section, top=T_MARGIN, bottom=B_MARGIN, left=L_MARGIN, right=R_MARGIN):
    section.top_margin = Emu(top * 914)   # twips to EMU: 1 twip = 914.4 EMU, approx
    section.bottom_margin = Emu(bottom * 914)
    section.left_margin = Emu(left * 914)
    section.right_margin = Emu(right * 914)


def body_para(doc, text='', size_pt=10, bold=False, italic=False, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    para.alignment = align
    if text:
        run = para.add_run(text)
        set_font(run, size_pt=size_pt, bold=bold, italic=italic, color=color)
    return para


def add_banner(doc, text, font_size=14):
    """Full-width navy title bar as a single-cell table row."""
    tbl = doc.add_table(rows=1, cols=1)
    remove_table_borders(tbl)
    set_table_width(tbl, CONTENT_W)
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, NAVY_HEX)
    set_row_height(tbl.rows[0], 400, exact=True)  # consistent height ~0.28in
    set_cell_margins(cell, top=60, bottom=60, left=115, right=115)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_font(r, size_pt=font_size, bold=True, color=WHITE)
    return tbl


def add_banner_continued(doc, text, font_size=13):
    """Sub-banner (continued pages) — 13pt bold white, navy bg."""
    return add_banner(doc, text, font_size=font_size)


def add_h3(doc, text):
    """H3 subhead — 11pt bold navy, 1pt CCCCCC bottom border."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    set_para_bottom_border(para, BORDER_HEX, 8)
    run = para.add_run(text)
    set_font(run, size_pt=11, bold=True, color=NAVY)
    return para


def add_section_divider(doc, text):
    """Section divider — 12pt bold navy, bottom border CCCCCC."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(8)
    set_para_bottom_border(para, BORDER_HEX, 8)
    run = para.add_run(text)
    set_font(run, size_pt=12, bold=True, color=NAVY)
    return para


def add_bullet(doc, text, lead_bold=None):
    """Bullet list item — 10pt, optional bold lead term."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.left_indent = Pt(18)
    # Bullet character
    run_bul = para.add_run('\u2022  ')
    set_font(run_bul, size_pt=10, color=GRAY)
    if lead_bold:
        run_lead = para.add_run(lead_bold)
        set_font(run_lead, size_pt=10, bold=True, color=GRAY)
        rest = text
        if rest:
            run_rest = para.add_run(rest)
            set_font(run_rest, size_pt=10, color=GRAY)
    else:
        run = para.add_run(text)
        set_font(run, size_pt=10, color=GRAY)
    return para


def add_check_bullet(doc, text):
    """Check-mark bullet (highlights box) — light blue row."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.left_indent = Pt(14)
    set_para_shading(para, LT_BLUE_HEX)
    run_check = para.add_run('\u2713  ')
    set_font(run_check, size_pt=10, bold=True, color=BLUE)
    run_text = para.add_run(text)
    set_font(run_text, size_pt=10, color=GRAY)
    return para


def insert_toc_field(doc):
    """Insert Word TOC field for LibreOffice to update."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    return para


def add_two_col_rec_table(doc, left_sections, right_sections):
    """Add a 2-column recommendations table (no visible borders).
    left_sections and right_sections: list of (title, bullets) tuples.
    """
    rows = max(len(left_sections), len(right_sections))
    tbl = doc.add_table(rows=rows, cols=2)
    remove_table_borders(tbl)
    set_table_width(tbl, CONTENT_W)
    half = CONTENT_W // 2

    for i in range(rows):
        row = tbl.rows[i]
        for j, sections in enumerate([left_sections, right_sections]):
            cell = row.cells[j]
            set_cell_width(cell, half)
            set_cell_margins(cell, top=40, bottom=40, left=60, right=60)
            set_no_cell_borders(cell)

            if i < len(sections):
                title, bullets = sections[i]
                # H3 in BLUE (10pt bold) — only when title is non-empty
                if title:
                    p = cell.add_paragraph()
                    p.paragraph_format.space_before = Pt(10)
                    p.paragraph_format.space_after = Pt(4)
                    set_para_bottom_border(p, BORDER_HEX, 4)
                    r = p.add_run(title)
                    set_font(r, size_pt=10, bold=True, color=BLUE)
                # Bullets
                for bullet in bullets:
                    pb = cell.add_paragraph()
                    pb.paragraph_format.space_before = Pt(2)
                    pb.paragraph_format.space_after = Pt(2)
                    pb.paragraph_format.left_indent = Pt(10)
                    rb = pb.add_run(f'\u2022  {bullet}')
                    set_font(rb, size_pt=9, color=GRAY)
    return tbl


# ─────────────────────────────────────────────────────────────────────────────
# HEADER / FOOTER BUILDERS
# ─────────────────────────────────────────────────────────────────────────────

def build_standard_header(section):
    """Standard header: logo left + doc-tag right + 2pt navy bottom rule."""
    header = section.header
    header.is_linked_to_previous = False

    # Clear any existing paragraphs
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)

    para = header.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(3)
    set_para_bottom_border(para, 'AAAAAA', 4)  # 0.5pt light gray rule

    # Logo (if exists)
    if os.path.exists(LOGO_H_PATH):
        run_logo = para.add_run()
        run_logo.add_picture(LOGO_H_PATH, height=Pt(28))
    else:
        run_logo = para.add_run('NBAIS')
        set_font(run_logo, size_pt=12, bold=True, color=NAVY)

    # Tab
    para.add_run('\t')

    # Doc-tag
    run_tag = para.add_run("Workers\u2019 Compensation Proposal")
    set_font(run_tag, size_pt=9, italic=True, color=GRAY)

    # Right-align tab stop
    add_right_tab_stop(para, CONTENT_W)


def build_standard_footer(section, right_label):
    """Standard footer: left text + right label + 1pt mid-gray top rule."""
    footer = section.footer
    footer.is_linked_to_previous = False

    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)

    para = footer.add_paragraph()
    set_para_top_border(para, MID_GRAY_HEX, 8)  # 1pt mid-gray
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(0)

    # Left text ({memberName} is a docxtemplater tag — literal in template)
    run_left = para.add_run("NBAIS Workers\u2019 Compensation Proposal \u00b7 {memberName} \u00b7 Confidential")
    set_font(run_left, size_pt=7.5, color=GRAY)

    # Tab
    para.add_run('\t')

    # Right label
    run_right = para.add_run(right_label)
    set_font(run_right, size_pt=7.5, color=GRAY)

    add_right_tab_stop(para, CONTENT_W)


def link_header(section):
    """Link section header to previous section."""
    section.header.is_linked_to_previous = True


def apply_standard_margins(section):
    """Apply standard page size and margins to section (1-inch all sides for 6.5in text width)."""
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    section.different_first_page_header_footer = False


# ─────────────────────────────────────────────────────────────────────────────
# KV TABLE HELPER
# ─────────────────────────────────────────────────────────────────────────────

def add_kv_table(doc, rows_data, label_pct=35, banner_text=None):
    """2-column label|value table with optional navy banner first row.
    rows_data: list of (label, value) tuples. None separator → light-gray row spanning full width.
    """
    label_w = int(CONTENT_W * label_pct / 100)
    value_w = CONTENT_W - label_w

    n_rows = len(rows_data) + (1 if banner_text else 0)
    tbl = doc.add_table(rows=n_rows, cols=2)
    set_table_borders(tbl, color=BORDER_HEX, size=4)
    set_table_width(tbl, CONTENT_W)

    row_idx = 0
    if banner_text:
        banner_row = tbl.rows[row_idx]
        # Merge cells across the 2 cols
        cell0 = banner_row.cells[0]
        cell1 = banner_row.cells[1]
        cell0.merge(cell1)
        set_cell_bg(cell0, NAVY_HEX)
        set_cell_margins(cell0, top=80, bottom=80, left=100, right=80)
        p = cell0.paragraphs[0]
        r = p.add_run(banner_text)
        set_font(r, size_pt=11, bold=True, color=WHITE)
        row_idx += 1

    for i, (label, value) in enumerate(rows_data):
        tr = tbl.rows[row_idx]
        row_idx += 1

        lc = tr.cells[0]
        vc = tr.cells[1]
        set_cell_width(lc, label_w)
        set_cell_width(vc, value_w)
        set_cell_margins(lc, top=80, bottom=80, left=115, right=115)
        set_cell_margins(vc, top=80, bottom=80, left=115, right=115)
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Alternating row background
        if i % 2 == 0:
            set_cell_bg(lc, LT_GRAY_HEX)
            set_cell_bg(vc, LT_GRAY_HEX)

        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        set_font(lr, size_pt=10, bold=True, color=NAVY)

        vp = vc.paragraphs[0]
        vr = vp.add_run(str(value))
        set_font(vr, size_pt=10, color=GRAY)

    return tbl


# ─────────────────────────────────────────────────────────────────────────────
# PAGE 1 — COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────

def build_cover_page(doc):
    """Build cover page content (Section 1, no header/footer)."""

    # 1. Full-width navy rule (edge-to-edge since margins = 0)
    tbl_rule = doc.add_table(rows=1, cols=1)
    remove_table_borders(tbl_rule)
    set_table_width(tbl_rule, PAGE_W)
    cell_rule = tbl_rule.rows[0].cells[0]
    set_cell_bg(cell_rule, NAVY_HEX)
    set_row_height(tbl_rule.rows[0], 460, exact=True)  # ~0.32in
    set_cell_margins(cell_rule, top=0, bottom=0, left=0, right=0)
    p = cell_rule.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # 2. Eyebrow text
    p_eyebrow = doc.add_paragraph()
    p_eyebrow.paragraph_format.space_before = Pt(14)
    p_eyebrow.paragraph_format.space_after = Pt(4)
    p_eyebrow.paragraph_format.left_indent = Inches(0.55)
    r = p_eyebrow.add_run('Nevada Builders Alliance Insurance Solutions')
    set_font(r, size_pt=9, italic=True, color=GRAY)

    # 3. Stacked logo placeholder (docxtemplater image tag)
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(20)
    p_logo.paragraph_format.space_after = Pt(20)
    r_logo = p_logo.add_run('{%stackedLogoBase64}')
    set_font(r_logo, size_pt=10, color=NAVY)

    # 4. Cover title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(8)
    r_title1 = p_title.add_run("Workers\u2019 Compensation")
    set_font(r_title1, size_pt=26, bold=True, color=NAVY)
    p_title.add_run().add_break()  # line break within paragraph
    r_title2 = p_title.add_run("Insurance Proposal")
    set_font(r_title2, size_pt=26, bold=True, color=NAVY)

    # 5. Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(30)
    r_sub = p_sub.add_run('Prepared exclusively for Nevada Builders Alliance members')
    set_font(r_sub, size_pt=11, italic=True, color=GRAY)

    # 6. Meta grid table (centered, 2-col, no borders)
    meta_rows = [
        ('Prepared For',  '{memberName}'),
        ('Policy Period', '{policyPeriod}'),
        ('Prepared By',   'Dianne Slater'),
        ('Date',          '{quoteDate}'),
        ('Program',       'Nevada Builders Alliance \u2014 NBAIS Member Program'),
    ]
    label_w = int(CONTENT_W * 0.35)
    value_w = CONTENT_W - label_w

    tbl_meta = doc.add_table(rows=len(meta_rows), cols=2)
    remove_table_borders(tbl_meta)
    set_table_width(tbl_meta, CONTENT_W)
    # Center the table by adding left indent equivalent to left margin
    tbl_meta.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(meta_rows):
        row = tbl_meta.rows[i]
        lc = row.cells[0]
        vc = row.cells[1]
        set_cell_width(lc, label_w)
        set_cell_width(vc, value_w)
        set_cell_margins(lc, top=40, bottom=40, left=80, right=60)
        set_cell_margins(vc, top=40, bottom=40, left=60, right=80)

        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lr = lp.add_run(label)
        set_font(lr, size_pt=10, bold=True, color=NAVY)

        vp = vc.paragraphs[0]
        vr = vp.add_run(value)
        set_font(vr, size_pt=10, color=GRAY)

    # 7. Spacer + cover footer
    for _ in range(4):
        ps = doc.add_paragraph()
        ps.paragraph_format.space_before = Pt(0)
        ps.paragraph_format.space_after = Pt(0)

    p_cfooter = doc.add_paragraph()
    p_cfooter.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cfooter.paragraph_format.space_before = Pt(20)
    p_cfooter.paragraph_format.space_after = Pt(0)
    set_para_top_border(p_cfooter, NAVY_HEX, 12)  # 1.5pt navy
    r_cf = p_cfooter.add_run(
        'Confidential \u2014 Prepared for the named member\u2019s exclusive use'
    )
    set_font(r_cf, size_pt=8.5, italic=True, color=GRAY)


# ─────────────────────────────────────────────────────────────────────────────
# PAGE 2 — TABLE OF CONTENTS
# ─────────────────────────────────────────────────────────────────────────────

# TODO: Re-enable TOC when heading styles are applied to section headings
def build_toc_page(doc):
    """Build TOC page content (Section 2)."""
    add_banner(doc, 'Table of Contents', font_size=14)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(8)
    p_space.paragraph_format.space_after = Pt(0)

    insert_toc_field(doc)

    # Placeholder paragraphs (TOC expands on update)
    for _ in range(12):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)


# ─────────────────────────────────────────────────────────────────────────────
# PAGE 3 — COVER LETTER
# ─────────────────────────────────────────────────────────────────────────────

def build_cover_letter_page(doc):
    """Build cover letter page content (Section 3)."""

    # Letter meta block
    for tag in ['{quoteDate}', '{memberName}', '{memberAddress}']:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(tag)
        set_font(r, size_pt=10, color=GRAY)

    # RE: line
    p_re = doc.add_paragraph()
    p_re.paragraph_format.space_before = Pt(10)
    p_re.paragraph_format.space_after = Pt(10)
    r_re = p_re.add_run(
        'RE: Workers\u2019 Compensation Insurance Proposal \u2014 Nevada Builders Alliance Member Program'
    )
    set_font(r_re, size_pt=10, bold=True, color=GRAY)

    # Salutation
    p_sal = body_para(doc, 'Dear {memberName},', size_pt=10, color=GRAY, space_before=0, space_after=10)

    # About section
    add_h3(doc, 'About this proposal')

    body_para(doc,
        'On behalf of Nevada Builders Alliance Insurance Services (NBAIS), we are pleased to present '
        'this Workers\u2019 Compensation Insurance proposal exclusively for members of the Nevada '
        'Builders Alliance (NBA). This proposal has been prepared specifically for your organization '
        'and reflects the competitive program rates and enhanced coverage options available through '
        'your NBA membership.',
        size_pt=10, color=GRAY, space_before=0, space_after=6)

    body_para(doc,
        'NBAIS was established to serve the unique risk management needs of Nevada\u2019s construction '
        'industry \u2014 from residential and commercial builders to specialty trade contractors. As an '
        'NBA member, your organization has access to a Workers\u2019 Compensation program designed '
        'around the realities of your trade, not a one-size-fits-all solution.',
        size_pt=10, color=GRAY, space_before=0, space_after=6)

    # Program highlights
    add_h3(doc, 'Program highlights')

    highlights = [
        'Exclusive NBA member pricing \u2014 competitive group rates unavailable in the open market',
        'Construction-class expertise \u2014 underwriting specialists who understand your trade',
        'Dividend potential \u2014 SIG participation with return of premium for favorable loss performance',
        'Loss control resources \u2014 proactive safety and claims management support',
        'Dedicated service team \u2014 NBAIS producers with direct carrier access',
    ]
    for h in highlights:
        add_bullet(doc, h)

    # What is included
    add_h3(doc, 'What is included in this proposal')

    body_para(doc, 'This proposal package contains the following for your review:',
              size_pt=10, color=GRAY, space_before=0, space_after=6)

    add_bullet(doc, ' \u2014 a summary of your proposed coverage terms and estimated premium.',
               lead_bold='Premium Summary & Coverage at a Glance')
    add_bullet(doc, ' \u2014 a detailed outline of the proposed coverage terms, limits, and exclusions applicable to your operation.',
               lead_bold='Workers\u2019 Compensation Coverage Details')
    add_bullet(doc, ' \u2014 the carrier quotation secured for your review, including the class code and payroll basis used to develop this quote. Please review for accuracy and notify us of any changes prior to binding.',
               lead_bold='Carrier Quote')
    add_bullet(doc, ' \u2014 a comprehensive list of additional coverage lines for your consideration across commercial, personal, bond, employee benefits, and life planning categories.',
               lead_bold='Coverage Recommendations')


# ─────────────────────────────────────────────────────────────────────────────
# PAGE 4 — PREMIUM SUMMARY + COVERAGE DETAILS (consolidated)
# ─────────────────────────────────────────────────────────────────────────────

def build_premium_summary_page(doc):
    """Build premium summary + coverage details consolidated page (Section 4)."""
    add_banner(doc, 'Premium Summary', font_size=14)

    # Lead italic
    p_lead = doc.add_paragraph()
    p_lead.paragraph_format.space_before = Pt(8)
    p_lead.paragraph_format.space_after = Pt(10)
    r_lead = p_lead.add_run(
        'Your estimated cost for the coverage period {policyPeriod}. '
        'All figures are subject to final payroll audit.'
    )
    set_font(r_lead, size_pt=10, italic=True, color=GRAY)

    # Coverage at a Glance table
    cag_rows = [
        ('Insured',                   '{memberName}'),
        ('Policy Period',             '{policyPeriod}'),
        ('Coverage',                  'Workers\u2019 Compensation \u2014 Statutory (Nevada) / Employers\u2019 Liability'),
        ('Employers\u2019 Liability Limits',
         '$1,000,000 Each Accident / $1,000,000 Disease \u2013 Each Employee / $1,000,000 Disease \u2013 Policy Limit'),
        ('Program',                   'Nevada Builders Alliance \u2014 NBAIS Member Program'),
        ('Carrier',                   'Builders Association of Western Nevada Self-Insured Group (BAWNSIG)'),
        ('Est. Premium',              '{estPremium} (subject to final audit)'),
        ('Surplus Contribution (8%)', '{surplusContribution}'),
        ('Employers\u2019 Liability Fee',  '{employersLiabilityFee}'),
    ]

    label_w = int(CONTENT_W * 0.35)
    value_w = CONTENT_W - label_w

    tbl = doc.add_table(rows=len(cag_rows) + 3, cols=2)  # +1 banner, +1 total, +1 downpayment
    set_table_borders(tbl, color=BORDER_HEX, size=4)
    set_table_width(tbl, CONTENT_W)

    # Banner row
    b_cell = tbl.rows[0].cells[0]
    b_cell.merge(tbl.rows[0].cells[1])
    set_cell_bg(b_cell, NAVY_HEX)
    set_cell_margins(b_cell, top=80, bottom=80, left=100, right=80)
    p = b_cell.paragraphs[0]
    r = p.add_run('Coverage at a Glance')
    set_font(r, size_pt=11, bold=True, color=WHITE)
    set_row_header(tbl.rows[0])

    # Data rows
    for i, (label, value) in enumerate(cag_rows):
        row = tbl.rows[i + 1]
        lc, vc = row.cells[0], row.cells[1]
        set_cell_width(lc, label_w)
        set_cell_width(vc, value_w)
        set_cell_margins(lc, top=80, bottom=80, left=115, right=115)
        set_cell_margins(vc, top=80, bottom=80, left=115, right=115)
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if i % 2 == 0:
            set_cell_bg(lc, LT_GRAY_HEX)
            set_cell_bg(vc, LT_GRAY_HEX)
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        set_font(lr, size_pt=10, bold=True, color=NAVY)
        vp = vc.paragraphs[0]
        vr = vp.add_run(value)
        set_font(vr, size_pt=10, color=GRAY)

    # Total row
    total_row = tbl.rows[len(cag_rows) + 1]
    tlc, tvc = total_row.cells[0], total_row.cells[1]
    set_cell_bg(tlc, LT_BLUE_HEX)
    set_cell_bg(tvc, LT_BLUE_HEX)
    set_cell_margins(tlc, top=80, bottom=80, left=80, right=60)
    set_cell_margins(tvc, top=80, bottom=80, left=80, right=60)
    tp = tlc.paragraphs[0]
    tr = tp.add_run('Total Estimated Cost')
    set_font(tr, size_pt=11, bold=True, color=NAVY)
    tvp = tvc.paragraphs[0]
    tvr = tvp.add_run('{totalEstimatedPremium}')
    set_font(tvr, size_pt=11, bold=True, color=NAVY)

    # Down payment row
    dp_row = tbl.rows[len(cag_rows) + 2]
    dlc, dvc = dp_row.cells[0], dp_row.cells[1]
    set_cell_margins(dlc, top=60, bottom=60, left=80, right=60)
    set_cell_margins(dvc, top=60, bottom=60, left=80, right=60)
    dp = dlc.paragraphs[0]
    dr = dp.add_run('Initial Down Payment')
    set_font(dr, size_pt=10, bold=True, color=NAVY)
    dvp = dvc.paragraphs[0]
    dvr = dvp.add_run(
        '{downPayment} (25% \u2014 new business). Balance payable online via secure payment link provided upon binding.'
    )
    set_font(dvr, size_pt=10, color=GRAY)

    # What's next section
    add_h3(doc, "What\u2019s next")
    body_para(doc,
        'Review the Coverage Details on the following page, confirm payroll and class code accuracy, '
        'and contact your NBAIS producer to bind. Final premium will be reconciled at audit.',
        size_pt=10, color=GRAY, space_before=0, space_after=8)

    # Explicit page break — force Coverage Details to start on new page
    p_break = doc.add_paragraph()
    p_break.paragraph_format.space_before = Pt(0)
    p_break.paragraph_format.space_after = Pt(0)
    run_break = p_break.add_run()
    run_break.add_break(WD_BREAK.PAGE)

    # Coverage Details sub-section
    add_banner_continued(doc, 'Coverage Details \u2014 Workers\u2019 Compensation', font_size=13)

    add_h3(doc, 'Policy Information')

    pi_rows = [
        ('Carrier',             'Builders Association of Western Nevada Self-Insured Group (BAWNSIG)'),
        ('Program Manager',     'Lusense'),
        ('Financial Strength',  'BAWNSIG is a Nevada state-regulated self-insured group. AM Best rating not applicable \u2014 see program disclosure.'),
        ('Policy Period',       '{policyPeriod}'),
        ('Coverage',            'Workers\u2019 Compensation'),
        ('States Covered',      'Nevada'),
    ]
    add_kv_table(doc, pi_rows)

    add_h3(doc, 'Named Insured')
    p_ni = doc.add_paragraph()
    p_ni.paragraph_format.space_before = Pt(4)
    p_ni.paragraph_format.space_after = Pt(4)
    r_ni = p_ni.add_run('{memberLegalName}')
    set_font(r_ni, size_pt=11, bold=True, color=NAVY)

    add_h3(doc, 'Coverage and Limits')

    cov_label_w = int(CONTENT_W * 0.65)
    cov_value_w = CONTENT_W - cov_label_w

    cov_rows = [
        ('Part I \u2014 Workers\u2019 Compensation',                   'Statutory per State of Nevada'),
        ('Part II \u2014 Employers\u2019 Liability: Each Accident',    '$1,000,000'),
        ('Part II \u2014 Employers\u2019 Liability: Disease \u2013 Each Employee', '$1,000,000'),
        ('Part II \u2014 Employers\u2019 Liability: Disease \u2013 Policy Limit',  '$1,000,000'),
    ]

    tbl_cov = doc.add_table(rows=len(cov_rows) + 1, cols=2)
    set_table_borders(tbl_cov, color=BORDER_HEX, size=4)
    set_table_width(tbl_cov, CONTENT_W)

    # Header row
    hr = tbl_cov.rows[0]
    hlc, hvc = hr.cells[0], hr.cells[1]
    set_cell_bg(hlc, NAVY_HEX)
    set_cell_bg(hvc, NAVY_HEX)
    set_cell_margins(hlc, top=70, bottom=70, left=80, right=60)
    set_cell_margins(hvc, top=70, bottom=70, left=80, right=60)
    hp1 = hlc.paragraphs[0]
    hp1.add_run('Coverage').font.bold = True
    hp1.runs[0].font.color.rgb = WHITE
    hp1.runs[0].font.size = Pt(10)
    hp1.runs[0].font.name = FONT
    hp2 = hvc.paragraphs[0]
    hp2.add_run('Limit').font.bold = True
    hp2.runs[0].font.color.rgb = WHITE
    hp2.runs[0].font.size = Pt(10)
    hp2.runs[0].font.name = FONT
    set_row_header(tbl_cov.rows[0])

    for i, (label, value) in enumerate(cov_rows):
        row = tbl_cov.rows[i + 1]
        lc, vc = row.cells[0], row.cells[1]
        set_cell_width(lc, cov_label_w)
        set_cell_width(vc, cov_value_w)
        set_cell_margins(lc, top=60, bottom=60, left=80, right=60)
        set_cell_margins(vc, top=60, bottom=60, left=80, right=60)
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if i % 2 == 0:
            set_cell_bg(lc, LT_GRAY_HEX)
            set_cell_bg(vc, LT_GRAY_HEX)
        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        set_font(lr, size_pt=10, color=GRAY)
        vp = vc.paragraphs[0]
        vr = vp.add_run(value)
        set_font(vr, size_pt=10, color=GRAY)

    add_h3(doc, 'Surplus Contribution')
    p_surplus = doc.add_paragraph()
    p_surplus.paragraph_format.space_before = Pt(4)
    p_surplus.paragraph_format.space_after = Pt(4)
    r1 = p_surplus.add_run(
        'As a self-insured group (SIG), BAWNSIG requires a surplus contribution in addition to the estimated premium. '
        'This contribution \u2014 calculated at 8% of the estimated premium \u2014 is a regulatory requirement for SIG '
        'participation in Nevada and supports the financial reserves of the group. '
    )
    set_font(r1, size_pt=10, color=GRAY)
    r2 = p_surplus.add_run('It is not a fee retained by NBAIS or your producer.')
    set_font(r2, size_pt=10, bold=True, color=GRAY)


# ─────────────────────────────────────────────────────────────────────────────
# PAGE 5 — COVERAGE DETAILS CONTINUED
# ─────────────────────────────────────────────────────────────────────────────

def build_coverage_details_continued_page(doc):
    """Build coverage details continued page (Section 5)."""
    add_banner_continued(doc, 'Coverage Details (continued)', font_size=13)

    # Employee Classification Schedule
    add_h3(doc, 'Employee Classification Schedule')

    # Column widths (proportional, total = CONTENT_W)
    # State 7% | Class Code 12% | Description 28% | Payroll 21% | Rate 12% | Premium 20%
    col_pcts = [7, 12, 28, 21, 12, 20]
    col_ws = [int(CONTENT_W * p / 100) for p in col_pcts]
    # Adjust last col for rounding
    col_ws[-1] = CONTENT_W - sum(col_ws[:-1])

    col_headers = ['State', 'Class Code', 'Description', 'Est. Annual Payroll', 'Rate', 'Est. Premium']

    tbl_cs = doc.add_table(rows=3, cols=6)  # header + 1 loop row + total row
    set_table_borders(tbl_cs, color=BORDER_HEX, size=4)
    set_table_width(tbl_cs, CONTENT_W)

    # Header row (navy)
    hr = tbl_cs.rows[0]
    for j, (hdr, w) in enumerate(zip(col_headers, col_ws)):
        cell = hr.cells[j]
        set_cell_width(cell, w)
        set_cell_bg(cell, NAVY_HEX)
        set_cell_margins(cell, top=70, bottom=70, left=60, right=60)
        p = cell.paragraphs[0]
        r = p.add_run(hdr)
        set_font(r, size_pt=9, bold=True, color=WHITE)
    set_row_header(tbl_cs.rows[0])

    # Data row (docxtemplater loop)
    dr = tbl_cs.rows[1]
    for j, w in enumerate(col_ws):
        cell = dr.cells[j]
        set_cell_width(cell, w)
        set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Cell 0: loop start + state tag
    c0 = dr.cells[0]
    p0a = c0.paragraphs[0]
    r0a = p0a.add_run('{#classSchedule}')
    set_font(r0a, size_pt=10, color=GRAY)
    p0b = c0.add_paragraph()
    r0b = p0b.add_run('{state}')
    set_font(r0b, size_pt=10, color=GRAY)

    # Cell 1: class code
    c1 = dr.cells[1]
    r1 = c1.paragraphs[0].add_run('{classCode}')
    set_font(r1, size_pt=10, color=GRAY)

    # Cell 2: description
    c2 = dr.cells[2]
    r2 = c2.paragraphs[0].add_run('{classDescription}')
    set_font(r2, size_pt=10, color=GRAY)

    # Cell 3: payroll
    c3 = dr.cells[3]
    r3 = c3.paragraphs[0].add_run('{estAnnualPayroll}')
    set_font(r3, size_pt=10, color=GRAY)

    # Cell 4: rate
    c4 = dr.cells[4]
    r4 = c4.paragraphs[0].add_run('{rate}')
    set_font(r4, size_pt=10, color=GRAY)

    # Cell 5: premium + loop end
    c5 = dr.cells[5]
    p5a = c5.paragraphs[0]
    r5a = p5a.add_run('{classEstPremium}')
    set_font(r5a, size_pt=10, color=GRAY)
    p5b = c5.add_paragraph()
    r5b = p5b.add_run('{/classSchedule}')
    set_font(r5b, size_pt=10, color=GRAY)

    # Total row
    tr_row = tbl_cs.rows[2]
    for j, w in enumerate(col_ws):
        cell = tr_row.cells[j]
        set_cell_width(cell, w)
        set_cell_bg(cell, LT_GRAY_HEX)
        set_cell_margins(cell, top=70, bottom=70, left=60, right=60)

    # Merge cols 0-4 in total row for label
    tc_label = tr_row.cells[0]
    for j in range(1, 5):
        tc_label.merge(tr_row.cells[j])
    tp_label = tc_label.paragraphs[0]
    tp_label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tr_label = tp_label.add_run('Total Estimated Premium')
    set_font(tr_label, size_pt=10, bold=True, color=NAVY)

    tc_total = tr_row.cells[5]
    tp_total = tc_total.paragraphs[0]
    tr_total = tp_total.add_run('{totalEstimatedPremium}')
    set_font(tr_total, size_pt=10, bold=True, color=NAVY)

    # Excluded persons — conditional block
    # Use {#hasExcludedPersons} boolean conditional
    p_ep_start = doc.add_paragraph()
    p_ep_start.paragraph_format.space_before = Pt(0)
    p_ep_start.paragraph_format.space_after = Pt(0)
    r_ep_start = p_ep_start.add_run('{#hasExcludedPersons}')
    set_font(r_ep_start, size_pt=1, color=WHITE)

    add_h3(doc, 'Excluded Persons')

    body_para(doc,
        'The following individuals have elected to reject Workers\u2019 Compensation coverage '
        'pursuant to NRS 616B.612 by filing Form D-43 with the Nevada Division of Industrial Relations.',
        size_pt=10, color=GRAY, space_before=0, space_after=6)

    # Excluded persons table (with inner loop)
    ep_label_w = int(CONTENT_W * 0.5)
    ep_value_w = CONTENT_W - ep_label_w

    tbl_ep = doc.add_table(rows=2, cols=2)  # header + loop row
    set_table_borders(tbl_ep, color=BORDER_HEX, size=4)
    set_table_width(tbl_ep, CONTENT_W)

    # Header row
    ep_hr = tbl_ep.rows[0]
    ep_h0 = ep_hr.cells[0]
    ep_h1 = ep_hr.cells[1]
    set_cell_bg(ep_h0, NAVY_HEX)
    set_cell_bg(ep_h1, NAVY_HEX)
    set_cell_margins(ep_h0, top=70, bottom=70, left=80, right=60)
    set_cell_margins(ep_h1, top=70, bottom=70, left=80, right=60)
    ep_h0.paragraphs[0].add_run('Name').font.bold = True
    ep_h0.paragraphs[0].runs[0].font.color.rgb = WHITE
    ep_h0.paragraphs[0].runs[0].font.size = Pt(10)
    ep_h0.paragraphs[0].runs[0].font.name = FONT
    ep_h1.paragraphs[0].add_run('Election Form').font.bold = True
    ep_h1.paragraphs[0].runs[0].font.color.rgb = WHITE
    ep_h1.paragraphs[0].runs[0].font.size = Pt(10)
    ep_h1.paragraphs[0].runs[0].font.name = FONT
    set_row_header(tbl_ep.rows[0])

    # Loop row
    ep_dr = tbl_ep.rows[1]
    ep_d0 = ep_dr.cells[0]
    ep_d1 = ep_dr.cells[1]
    set_cell_width(ep_d0, ep_label_w)
    set_cell_width(ep_d1, ep_value_w)
    set_cell_margins(ep_d0, top=60, bottom=60, left=80, right=60)
    set_cell_margins(ep_d1, top=60, bottom=60, left=80, right=60)
    ep_d0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    ep_d1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Cell 0: loop start + name
    ep_p0a = ep_d0.paragraphs[0]
    ep_p0a.add_run('{#excludedPersons}')
    set_font(ep_p0a.runs[0], size_pt=10, color=GRAY)
    ep_p0b = ep_d0.add_paragraph()
    ep_p0b.add_run('{name}')
    set_font(ep_p0b.runs[0], size_pt=10, color=GRAY)

    # Cell 1: form + loop end
    ep_p1a = ep_d1.paragraphs[0]
    ep_p1a.add_run('Form D-43 \u2014 Election to Reject Coverage')
    set_font(ep_p1a.runs[0], size_pt=10, color=GRAY)
    ep_p1b = ep_d1.add_paragraph()
    ep_p1b.add_run('{/excludedPersons}')
    set_font(ep_p1b.runs[0], size_pt=10, color=GRAY)

    # Fine print
    body_para(doc,
        'Note: Officers or members electing to reject coverage must file Form D-43 with the Nevada '
        'Division of Industrial Relations. Rejection is effective upon filing. These individuals '
        'will not be covered under this policy for work-related injuries.',
        size_pt=8.5, italic=True, color=GRAY, space_before=4, space_after=4)

    # End of conditional block
    p_ep_end = doc.add_paragraph()
    p_ep_end.paragraph_format.space_before = Pt(0)
    p_ep_end.paragraph_format.space_after = Pt(0)
    r_ep_end = p_ep_end.add_run('{/hasExcludedPersons}')
    set_font(r_ep_end, size_pt=1, color=WHITE)

    # Self-Insured Group Disclosure
    add_h3(doc, 'Self-Insured Group Disclosure')

    # SIG Disclosure box: light-gray bg + blue left border
    tbl_disc = doc.add_table(rows=1, cols=1)
    remove_table_borders(tbl_disc)
    set_table_width(tbl_disc, CONTENT_W)
    disc_cell = tbl_disc.rows[0].cells[0]
    set_cell_bg(disc_cell, 'E8E8E8')  # light gray
    set_cell_margins(disc_cell, top=100, bottom=100, left=130, right=100)
    # Blue left border accent
    set_cell_border(disc_cell, {
        'left': {'val': 'single', 'sz': 24, 'color': BLUE_HEX},
        'top': {'val': 'none', 'sz': 0, 'color': 'auto'},
        'bottom': {'val': 'none', 'sz': 0, 'color': 'auto'},
        'right': {'val': 'none', 'sz': 0, 'color': 'auto'},
    })
    p_disc = disc_cell.paragraphs[0]
    r_disc = p_disc.add_run(
        'BAWNSIG is a Nevada-regulated self-insured group, not a traditional insurance carrier, '
        'and therefore does not carry an AM Best financial strength rating. BAWNSIG operates under '
        'the regulatory oversight of the Nevada Division of Industrial Relations and maintains reserves '
        'in accordance with state requirements. Members of NBAIS benefit from the group\u2019s long-standing '
        'solvency and claims-paying history as a construction industry SIG in Nevada.'
    )
    set_font(r_disc, size_pt=9.5, color=GRAY)


# ─────────────────────────────────────────────────────────────────────────────
# PAGE 6 — NEXT STEPS & AUTHORIZATION
# ─────────────────────────────────────────────────────────────────────────────

def build_next_steps_page(doc):
    """Build next steps & authorization page (Section 6)."""
    add_banner(doc, 'Next Steps & Member Authorization', font_size=14)

    body_para(doc,
        'To bind coverage or to discuss this proposal in further detail, please contact your NBAIS '
        'producer using the information below. Please review all coverage details carefully and confirm '
        'payroll and class code accuracy prior to binding, as final premium is subject to audit.',
        size_pt=10, color=GRAY, space_before=8, space_after=8)

    # Contact grid — no outer border, shaded cells, inner vertical divider only
    tbl_contact = doc.add_table(rows=1, cols=2)
    remove_table_borders(tbl_contact)
    set_table_width(tbl_contact, CONTENT_W)
    # Set only the inner vertical divider
    tblPr = tbl_contact._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    insideV = OxmlElement('w:insideV')
    insideV.set(qn('w:val'), 'single')
    insideV.set(qn('w:sz'), '4')
    insideV.set(qn('w:space'), '0')
    insideV.set(qn('w:color'), 'CCCCCC')
    tblBorders.append(insideV)
    tblPr.append(tblBorders)
    half_w = CONTENT_W // 2

    for j, (title, lines) in enumerate([
        ('Your NBAIS Producer', [
            ('bold', 'Dianne Slater'),
            ('italic', 'Account Manager'),
            ('italic', '(775) 555-0100'),
            ('italic', 'dslater@nbais.com'),
        ]),
        ('NBAIS Program Office', [
            ('normal', 'Nevada Builders Alliance Insurance Services'),
            ('italic', '1234 Builder\u2019s Way, Reno, NV 89501'),
            ('italic', 'www.nbais.com'),
        ]),
    ]):
        cell = tbl_contact.rows[0].cells[j]
        set_cell_width(cell, half_w)
        set_cell_bg(cell, 'E8E8E8')
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)

        pt = cell.paragraphs[0]
        rt = pt.add_run(title)
        set_font(rt, size_pt=10, bold=True, color=NAVY)

        for style, text in lines:
            pl = cell.add_paragraph()
            rl = pl.add_run(text)
            set_font(rl, size_pt=10,
                     bold=(style == 'bold'),
                     italic=(style == 'italic'),
                     color=GRAY)

    # Authorization section
    add_h3(doc, 'Member Authorization')

    body_para(doc,
        'By signing below, the undersigned acknowledges receipt of this Workers\u2019 Compensation '
        'Insurance proposal and authorizes Nevada Builders Alliance Insurance Services (NBAIS) to bind '
        'coverage as described herein, effective on the policy period stated above. The undersigned '
        'confirms that the payroll, classification codes, and excluded persons listed in this proposal '
        'are accurate to the best of their knowledge and understands that final premium is subject to '
        'audit. The required initial down payment will be remitted online via the secure payment link '
        'provided upon binding.',
        size_pt=10, color=GRAY, space_before=4, space_after=12)

    # Signature table
    sig_rows = [
        ('By', ''),
        ('Print Name', ''),
        ('Title', ''),
        ('Date', ''),
    ]
    label_w = int(CONTENT_W * 0.25)
    line_w = CONTENT_W - label_w

    tbl_sig = doc.add_table(rows=len(sig_rows), cols=2)
    remove_table_borders(tbl_sig)
    set_table_width(tbl_sig, CONTENT_W)

    for i, (label, _) in enumerate(sig_rows):
        row = tbl_sig.rows[i]
        lc = row.cells[0]
        vc = row.cells[1]
        set_cell_width(lc, label_w)
        set_cell_width(vc, line_w)
        set_cell_margins(lc, top=80, bottom=80, left=0, right=60)
        set_cell_margins(vc, top=80, bottom=80, left=60, right=0)

        lp = lc.paragraphs[0]
        lr = lp.add_run(label)
        set_font(lr, size_pt=10, bold=True, color=NAVY)

        vp = vc.paragraphs[0]
        set_cell_border(vc, {'bottom': {'val': 'single', 'sz': 6, 'color': BORDER_HEX}})
        vp.add_run('')  # empty underline

    # Fine print
    body_para(doc,
        'This proposal is not a binder or guarantee of coverage. All coverage is subject to '
        'underwriting approval, policy terms, conditions, and exclusions. Premium estimates are '
        'subject to final payroll audit. NBAIS is an insurance program administered on behalf of '
        'Nevada Builders Alliance members.',
        size_pt=8.5, italic=True, color=GRAY, space_before=16, space_after=4)


# ─────────────────────────────────────────────────────────────────────────────
# PAGE 7 — COVERAGE RECOMMENDATIONS (1 of 3)
# ─────────────────────────────────────────────────────────────────────────────

def build_recommendations_1_page(doc):
    """Build coverage recommendations page 1 of 3 (Section 7)."""
    add_banner(doc, 'Coverage Recommendations', font_size=14)

    p_lead = doc.add_paragraph()
    p_lead.paragraph_format.space_before = Pt(8)
    p_lead.paragraph_format.space_after = Pt(8)
    r_lead = p_lead.add_run(
        'The following list identifies common coverage areas for your consideration. Please review '
        'with your NBAIS producer to determine which lines are recommended, currently insured, or '
        'not applicable to your operation.'
    )
    set_font(r_lead, size_pt=9.5, italic=True, color=GRAY)

    add_section_divider(doc, 'Commercial Lines')

    left_secs = [
        ('Property Coverages', [
            'Building / Business Personal Property',
            'Business Income & Extra Expense',
            'Equipment Breakdown',
            'Inland Marine / Contractors Equipment',
            'Installation Floater',
            'Builders Risk',
        ]),
        ('Cyber / Identity Theft / Crime', [
            'Cyber Liability',
            'Data Breach / Privacy Liability',
            'Identity Theft Protection',
            'Commercial Crime / Employee Dishonesty',
        ]),
        ('Workers\u2019 Compensation Coverages', [
            'Workers\u2019 Compensation \u2014 Statutory',
            'Employers\u2019 Liability',
            'Stop Gap / Employers Liability (Monopolistic States)',
        ]),
        ('Directors & Officers / EPL / Fiduciary', [
            'Directors & Officers Liability',
            'Employment Practices Liability (EPLI)',
            'Fiduciary Liability',
        ]),
    ]
    right_secs = [
        ('Liability Coverages', [
            'Commercial General Liability',
            'Products & Completed Operations',
            'Contractual Liability',
            'Personal & Advertising Injury',
        ]),
        ('Automobile Coverage', [
            'Commercial Auto Liability',
            'Physical Damage (Comp & Collision)',
            'Hired & Non-Owned Auto',
            'Motor Truck Cargo',
        ]),
        ('Umbrella / Excess Liability', [
            'Commercial Umbrella',
            'Excess Liability',
        ]),
        ('Errors & Omissions / Professional', [
            'Professional Liability / E&O',
            'Contractors Professional Liability',
            'Design-Build Professional Liability',
        ]),
    ]

    add_two_col_rec_table(doc, left_secs, right_secs)


# ─────────────────────────────────────────────────────────────────────────────
# PAGE 8 — COVERAGE RECOMMENDATIONS (2 of 3)
# ─────────────────────────────────────────────────────────────────────────────

def build_recommendations_2_page(doc):
    """Build coverage recommendations page 2 of 3 (Section 8)."""
    add_banner_continued(doc, 'Coverage Recommendations (continued)', font_size=13)

    add_section_divider(doc, 'Commercial Lines (continued)')

    left_secs = [
        ('Wind / Hail, Earthquake, Flood', [
            'Wind & Hail Coverage',
            'Earthquake Coverage',
            'Flood Coverage (NFIP or Private)',
        ]),
        ('Pollution Liability', [
            'Contractors Pollution Liability',
            'Environmental Impairment Liability',
            'Site Pollution Legal Liability',
        ]),
    ]
    right_secs = [
        ('Foreign Coverages', [
            'Foreign Voluntary Workers\u2019 Compensation',
            'Foreign General Liability',
            'Foreign Auto',
            'Foreign Package Policy',
        ]),
        ('', []),  # empty right cell for this row
    ]
    add_two_col_rec_table(doc, left_secs, right_secs)

    add_section_divider(doc, 'Personal Lines')

    p_left = [
        ('Personal Insurance', [
            'Automobile',
            'Home / Homeowners',
            'Flood / Earthquake',
            'Personal Umbrella',
        ]),
    ]
    p_right = [
        ('', [
            'Farm & Ranch',
            'Watercraft / Recreational Vehicles',
            'Personal Articles Floater',
        ]),
    ]
    add_two_col_rec_table(doc, p_left, p_right)

    add_section_divider(doc, 'Bond Recommendations')

    b_left = [
        ('Surety & Bonds', [
            'Contract Bond',
            'Court Bond',
            'Fidelity Bond',
            'Financial Institution Bond',
        ]),
    ]
    b_right = [
        ('', [
            'License & Permit Bond',
            'Probate Bond',
            'Public Official Bond',
            'Surety Bond',
        ]),
    ]
    add_two_col_rec_table(doc, b_left, b_right)


# ─────────────────────────────────────────────────────────────────────────────
# PAGE 9 — EMPLOYEE BENEFITS RECOMMENDATIONS (3 of 3)
# ─────────────────────────────────────────────────────────────────────────────

def build_employee_benefits_page(doc):
    """Build employee benefits recommendations page (Section 9)."""
    add_banner(doc, 'Employee Benefits Recommendations', font_size=14)

    body_para(doc,
        'Group benefits, life planning, and retirement plan services available through NBAIS '
        'for member consideration.',
        size_pt=9.5, italic=True, color=GRAY, space_before=8, space_after=8)

    add_section_divider(doc, 'Group Benefits')

    gb_left = [
        ('Health & Welfare', [
            'HR Services',
            'Group Medical',
            'Group Dental',
            'Vision',
            'Group Life and Accidental Death & Dismemberment (AD&D)',
        ]),
    ]
    gb_right = [
        ('Disability & Supplemental', [
            'Long Term Care',
            'Short Term Disability',
            'Section 125 Cafeteria Plans',
            'Individual Medical / Dental',
        ]),
    ]
    add_two_col_rec_table(doc, gb_left, gb_right)

    add_section_divider(doc, 'Life Department')

    life_left = [
        ('', [
            'Business Planning',
            'Estate Planning',
        ]),
    ]
    life_right = [
        ('', []),
    ]
    add_two_col_rec_table(doc, life_left, life_right)

    add_section_divider(doc, 'Retirement Plan Services')

    ret_left = [
        ('', [
            'Qualified Plans',
            'Non-Qualified Plans',
        ]),
    ]
    ret_right = [
        ('', []),
    ]
    add_two_col_rec_table(doc, ret_left, ret_right)

    # Callout box
    p_callout = doc.add_paragraph()
    p_callout.paragraph_format.space_before = Pt(14)
    p_callout.paragraph_format.space_after = Pt(4)
    p_callout.paragraph_format.left_indent = Pt(14)
    p_callout.paragraph_format.right_indent = Pt(14)
    set_para_shading(p_callout, LT_BLUE_HEX)
    r_bold = p_callout.add_run('Discuss with your producer. ')
    set_font(r_bold, size_pt=9, bold=True, color=NAVY)
    r_callout = p_callout.add_run(
        'Your NBAIS producer can help you assess which of these coverage lines apply to your '
        'operation and identify any potential gaps in your current insurance program.'
    )
    set_font(r_callout, size_pt=9, color=NAVY)


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

def main():
    os.makedirs(OUTDIR, exist_ok=True)
    doc = Document()

    # Remove default empty paragraph
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    # ── Section 1: Cover page (no header/footer) ──────────────────────────
    s1 = doc.sections[0]
    s1.page_width    = Inches(8.5)
    s1.page_height   = Inches(11)
    s1.top_margin    = Inches(0)
    s1.bottom_margin = Inches(0)
    s1.left_margin   = Inches(0)
    s1.right_margin  = Inches(0)
    s1.different_first_page_header_footer = True
    # First-page header and footer are empty (cover has none)
    s1.first_page_header.is_linked_to_previous = False
    s1.first_page_footer.is_linked_to_previous = False

    build_cover_page(doc)

    # ── Section 3: Cover Letter ───────────────────────────────────────────
    s3 = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_standard_margins(s3)
    build_standard_header(s3)  # First content section 2014 build header here (TOC removed)
    build_standard_footer(s3, 'Cover Letter')
    build_cover_letter_page(doc)

    # ── Section 4: Premium Summary + Coverage Details ─────────────────────
    s4 = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_standard_margins(s4)
    link_header(s4)
    build_standard_footer(s4, 'Premium Summary')
    build_premium_summary_page(doc)

    # ── Section 5: Coverage Details continued ─────────────────────────────
    s5 = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_standard_margins(s5)
    link_header(s5)
    build_standard_footer(s5, 'Coverage Details (2 of 2)')
    build_coverage_details_continued_page(doc)

    # ── Section 6: Next Steps & Authorization ────────────────────────────
    s6 = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_standard_margins(s6)
    link_header(s6)
    build_standard_footer(s6, 'Next Steps & Authorization')
    build_next_steps_page(doc)

    # ── Section 7: Coverage Recommendations (1 of 3) ─────────────────────
    s7 = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_standard_margins(s7)
    link_header(s7)
    build_standard_footer(s7, 'Coverage Recommendations (1 of 3)')
    build_recommendations_1_page(doc)

    # ── Section 8: Coverage Recommendations (2 of 3) ─────────────────────
    s8 = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_standard_margins(s8)
    link_header(s8)
    build_standard_footer(s8, 'Coverage Recommendations (2 of 3)')
    build_recommendations_2_page(doc)

    # ── Section 9: Employee Benefits Recommendations ──────────────────────
    s9 = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_standard_margins(s9)
    link_header(s9)
    build_standard_footer(s9, 'Coverage Recommendations (3 of 3)')
    build_employee_benefits_page(doc)

    # ── Save ──────────────────────────────────────────────────────────────
    doc.save(OUTPATH)
    print(f'Saved: {OUTPATH}')


def parse_args():
    parser = argparse.ArgumentParser(description='Build NBAIS WC master.docx template.')
    parser.add_argument(
        '--sync',
        action='store_true',
        help='After generating master.docx, sync templates/verticals/nbais-wc/ to S3.',
    )
    return parser.parse_args()


if __name__ == '__main__':
    args = parse_args()
    main()
    if args.sync:
        sync_cmd = [
            'aws', 's3', 'sync',
            'templates/verticals/nbais-wc/',
            's3://fortress-tools/fip-proposal-templates/verticals/nbais-wc/',
            '--profile', 'fortress-tools-deployer',
            '--region', 'us-east-1',
            '--exact-timestamps',
        ]
        print(f'Syncing to S3: {" ".join(sync_cmd)}')
        subprocess.run(sync_cmd, check=True)
