#!/usr/bin/env python3
"""
Generate all docxtemplater-compatible proposal templates.
Run from: /home/fredw/projects/fip/services/proposal-generator
Usage: python3 scripts/generate-templates.py
"""

import os
import json
import shutil
import zipfile
import tempfile
from pathlib import Path
from datetime import datetime, timezone

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is required. Install with: pip install python-docx")
    raise SystemExit(1)


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def set_keep_with_next(paragraph):
    """Set keepNext on paragraph - prevents orphan rows in tables."""
    pPr = paragraph._p.get_or_add_pPr()
    keepNext = OxmlElement('w:keepNext')
    pPr.append(keepNext)


def set_repeat_header(row):
    """Mark row as repeat header (tblHeader)."""
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    trPr.append(tblHeader)


def set_keep_lines(cell):
    """Set keepLines on all paragraphs in cell."""
    for paragraph in cell.paragraphs:
        pPr = paragraph._p.get_or_add_pPr()
        keepLines = OxmlElement('w:keepLines')
        pPr.append(keepLines)


def set_cell_shading(cell, fill_color):
    """Apply background color to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)


HEADER_BG = 'BDD7EE'   # light blue
CURRENCY_COLS = {'premium', 'total', 'payroll', 'rate', 'taxes', 'fees', 'cost', 'subtotal'}


def apply_table_formatting(table, has_header=True, header_bg=HEADER_BG):
    """Apply all required table formatting: header shading/bold, keepNext, keepLines."""
    for i, row in enumerate(table.rows):
        is_header = has_header and i == 0
        is_last = i == len(table.rows) - 1
        if is_header:
            set_repeat_header(row)
        for j, cell in enumerate(row.cells):
            set_keep_lines(cell)
            if is_header and header_bg:
                set_cell_shading(cell, header_bg)
            for para in cell.paragraphs:
                if not is_last:
                    set_keep_with_next(para)
                if is_header:
                    for run in para.runs:
                        run.bold = True
                    # If cell was set via .text = '...' (no explicit runs), bold the implicit run
                    if not para.runs and para.text:
                        run = para.runs[0] if para.runs else para.add_run(para.text)
                        run.bold = True


def inject_update_fields(docx_path):
    """Inject updateFields setting so Word auto-updates TOC on open."""
    tmp = docx_path + '.tmp'
    with zipfile.ZipFile(docx_path, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/settings.xml':
                    data_str = data.decode('utf-8')
                    if 'w:updateFields' not in data_str:
                        data_str = data_str.replace(
                            '</w:settings>',
                            '<w:updateFields w:val="true"/></w:settings>'
                        )
                    data = data_str.encode('utf-8')
                zout.writestr(item, data)
    shutil.move(tmp, docx_path)


def add_schedule_section(doc, heading_text, heading_level=2, include_children=True):
    """Add a schedule items section with optional children loop."""
    doc.add_heading(heading_text, level=heading_level)

    p_sched_open = doc.add_paragraph('{#scheduleItems}')

    p_item_hdr = doc.add_paragraph()
    p_item_hdr.style = doc.styles['Heading 3']
    p_item_hdr.add_run('Item {itemNumber}: {description}')

    # formattedAttributes sub-loop table
    fa_table = doc.add_table(rows=2, cols=2)
    fa_table.style = 'Table Grid'
    fa_table.rows[0].cells[0].text = 'Field'
    fa_table.rows[0].cells[1].text = 'Value'
    set_repeat_header(fa_table.rows[0])
    fa_row = fa_table.rows[1]
    fa_row.cells[0].paragraphs[0].clear()
    fa_row.cells[0].paragraphs[0].add_run('{#formattedAttributes}{label}')
    fa_row.cells[1].paragraphs[0].clear()
    fa_row.cells[1].paragraphs[0].add_run('{formattedValue}{/formattedAttributes}')
    apply_table_formatting(fa_table)

    if include_children:
        p_child_open = doc.add_paragraph('{#children}')

        p_child_hdr = doc.add_paragraph()
        p_child_hdr.style = doc.styles['Heading 4']
        p_child_hdr.add_run('Sub-item {itemNumber}: {description}')

        child_fa_table = doc.add_table(rows=2, cols=2)
        child_fa_table.style = 'Table Grid'
        child_fa_table.rows[0].cells[0].text = 'Field'
        child_fa_table.rows[0].cells[1].text = 'Value'
        set_repeat_header(child_fa_table.rows[0])
        c_fa_row = child_fa_table.rows[1]
        c_fa_row.cells[0].paragraphs[0].clear()
        c_fa_row.cells[0].paragraphs[0].add_run('{#formattedAttributes}{label}')
        c_fa_row.cells[1].paragraphs[0].clear()
        c_fa_row.cells[1].paragraphs[0].add_run('{formattedValue}{/formattedAttributes}')
        apply_table_formatting(child_fa_table)

        doc.add_paragraph('{/children}')

    doc.add_paragraph('{/scheduleItems}')


def add_carrier_info_table(doc, rows_data):
    """Add a carrier info table with label/value rows."""
    carrier_table = doc.add_table(rows=len(rows_data), cols=2)
    carrier_table.style = 'Table Grid'
    for i, (label, value) in enumerate(rows_data):
        carrier_table.rows[i].cells[0].text = label
        set_cell_shading(carrier_table.rows[i].cells[0], HEADER_BG)
        # Bold the label
        for para in carrier_table.rows[i].cells[0].paragraphs:
            for run in para.runs:
                run.bold = True
        carrier_table.rows[i].cells[1].paragraphs[0].clear()
        carrier_table.rows[i].cells[1].paragraphs[0].add_run(value)
    apply_table_formatting(carrier_table, has_header=False)
    return carrier_table


def add_endorsements_table(doc):
    """Add the standard endorsements section."""
    doc.add_heading('Endorsements', level=2)
    end_table = doc.add_table(rows=2, cols=2)
    end_table.style = 'Table Grid'
    end_table.rows[0].cells[0].text = 'Form #'
    end_table.rows[0].cells[1].text = 'Endorsement'
    set_repeat_header(end_table.rows[0])
    end_row = end_table.rows[1]
    end_row.cells[0].paragraphs[0].clear()
    end_row.cells[0].paragraphs[0].add_run('{#endorsements}{formNumber}')
    end_row.cells[1].paragraphs[0].clear()
    end_row.cells[1].paragraphs[0].add_run('{name}{/endorsements}')
    apply_table_formatting(end_table)
    return end_table


def add_notes_section(doc):
    """Add the conditional notes paragraph."""
    p_notes = doc.add_paragraph()
    p_notes.add_run('{#notes}')
    p_notes_content = doc.add_paragraph()
    run = p_notes_content.add_run('Notes: ')
    run.bold = True
    p_notes_content.add_run('{notes}')
    doc.add_paragraph('{/notes}')


def set_cell_color(cell, hex_color):
    """Set explicit font color on all runs in a cell."""
    for para in cell.paragraphs:
        for run in para.runs:
            rPr = run._r.get_or_add_rPr()
            color_el = OxmlElement('w:color')
            color_el.set(qn('w:val'), hex_color)
            rPr.append(color_el)


def add_section_banner(doc, title):
    """Add a full-width dark blue banner with white bold uppercase text."""
    section = doc.sections[0]
    # page_width/margins are Length objects (EMU); arithmetic returns raw EMU int
    text_width_emu = section.page_width - section.left_margin - section.right_margin
    # 1 twip = 635 EMU
    width_twips = str(int(text_width_emu // 635))

    banner_table = doc.add_table(rows=1, cols=1)
    banner_table.style = 'Table Grid'

    # Set table width to full text width via raw XML
    tbl = banner_table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), width_twips)
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    cell = banner_table.rows[0].cells[0]
    set_cell_shading(cell, '1F3864')

    # Set cell width too
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), width_twips)
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

    # Clear and set paragraph
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)


def add_carrier_info_paragraphs(doc, fields=None):
    """Add carrier info as bold label + value paragraph pairs (no table)."""
    if fields is None:
        fields = [
            ('CARRIER', '{carrier.name}'),
            ('A.M. BEST RATING', '{carrier.amBestRating}'),
            ('POLICY PERIOD', '{effectiveDate} \u2014 {expirationDate}'),
            ('COVERAGE', '{sectionTitle}'),
        ]
    for label, value in fields:
        p = doc.add_paragraph()
        run_label = p.add_run(label + '    ')
        run_label.bold = True
        run_label.font.size = Pt(10)
        run_value = p.add_run(value)
        run_value.bold = False
        run_value.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)


def add_named_insured_section(doc):
    """Add a Named Insured bullet list section using docxtemplater loop."""
    p_heading = doc.add_paragraph()
    p_heading.style = doc.styles['Heading 3']
    p_heading.add_run('Named Insured')

    p_open = doc.add_paragraph()
    p_open.add_run('{#namedInsureds}')

    p_item = doc.add_paragraph(style='List Bullet')
    p_item.add_run('\u2022 {name}')

    p_close = doc.add_paragraph()
    p_close.add_run('{/namedInsureds}')


# ---------------------------------------------------------------------------
# Master template
# ---------------------------------------------------------------------------

def create_master_docx(output_path):
    """Create the master proposal template."""
    doc = Document()

    # Remove default empty paragraph if present
    if doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    # --- Section 1: Cover Page ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style = doc.styles['Heading 1']
    run = p.add_run('Insurance Proposal')
    run.bold = True
    run.font.size = Pt(28)

    p2 = doc.add_paragraph('Prepared for:')

    p3 = doc.add_paragraph()
    run3 = p3.add_run('{insuredName}')
    run3.bold = True
    run3.font.size = Pt(16)

    p4 = doc.add_paragraph()
    p4.add_run('Policy Period: ')
    p4.add_run('{policyPeriodDisplay}')

    p5 = doc.add_paragraph()
    p5.add_run('Prepared by: ')
    p5.add_run('{amName}')
    p6 = doc.add_paragraph()
    p6.add_run('{amEmail}')

    p7 = doc.add_paragraph()
    p7.add_run('Date: ')
    p7.add_run('{generatedDate}')

    doc.add_paragraph('NBA Insurance Services')
    doc.add_page_break()

    # --- Section 2: Table of Contents (manual — no Word field codes) ---
    doc.add_heading('Table of Contents', level=1)

    toc_entries = [
        ('Your Service Team', False),
        ('Premium Summary', False),
        ('Market Response', False),
        ('Executive Summary', False),
        ('Coverage Details', False),
        ('    General Liability', True),
        ("    Workers' Compensation", True),
        ('    Commercial Property', True),
        ('Recommendations', False),
        ('About NBA Insurance Services', False),
    ]
    for entry_text, is_sub in toc_entries:
        p_toc = doc.add_paragraph()
        p_toc.paragraph_format.left_indent = Pt(18) if is_sub else Pt(0)
        # Compute dot leaders to fill ~55 chars total width (adjust based on indent)
        max_width = 50 if not is_sub else 46
        dots_needed = max(5, max_width - len(entry_text.strip()))
        dot_leaders = ' ' + ('.' * dots_needed) + ' \u2014'
        run_toc = p_toc.add_run(entry_text.strip() + dot_leaders)
        run_toc.font.size = Pt(11)

    doc.add_page_break()

    # --- Section 3: Team Sheet (conditional) ---
    p_team_open = doc.add_paragraph()
    p_team_open.add_run('{#hasTeam}')

    doc.add_heading('Your Service Team', level=1)

    team_table = doc.add_table(rows=2, cols=3)
    team_table.style = 'Table Grid'

    # Header row
    hdr_cells = team_table.rows[0].cells
    hdr_cells[0].text = 'Name / Role'
    hdr_cells[1].text = 'Email'
    hdr_cells[2].text = 'Phone'
    set_repeat_header(team_table.rows[0])
    for cell in team_table.rows[0].cells:
        set_keep_lines(cell)
        for para in cell.paragraphs:
            set_keep_with_next(para)

    # Template row — docxtemplater repeats this row for each team member
    row = team_table.rows[1]
    c0 = row.cells[0]
    c0.paragraphs[0].clear()
    c0.paragraphs[0].add_run('{#team}{name}')
    c0.add_paragraph().add_run('{role}')
    c1 = row.cells[1]
    c1.paragraphs[0].clear()
    c1.paragraphs[0].add_run('{email}')
    c2 = row.cells[2]
    c2.paragraphs[0].clear()
    c2.paragraphs[0].add_run('{phone}{/team}')

    for cell in row.cells:
        set_keep_lines(cell)

    doc.add_page_break()

    p_team_close = doc.add_paragraph()
    p_team_close.add_run('{/hasTeam}')

    # --- Section 4: Premium Summary ---
    add_section_banner(doc, 'PREMIUM SUMMARY')

    # Single unified premium table
    ps_table = doc.add_table(rows=2, cols=3)
    ps_table.style = 'Table Grid'

    hdr = ps_table.rows[0]
    hdr.cells[0].text = 'COVERAGE'
    hdr.cells[1].text = 'EXPOSURE / HIGHLIGHTS'
    hdr.cells[2].text = 'PREMIUM'
    set_repeat_header(hdr)

    # Template row — docxtemplater repeats for each premiumRows item
    data_row = ps_table.rows[1]
    data_row.cells[0].paragraphs[0].clear()
    data_row.cells[0].paragraphs[0].add_run('{#premiumRows}{coverageLabel}')
    data_row.cells[1].paragraphs[0].clear()
    data_row.cells[1].paragraphs[0].add_run('{exposureHighlights}')
    data_row.cells[2].paragraphs[0].clear()
    data_row.cells[2].paragraphs[0].add_run('{formattedPremium}{/premiumRows}')
    apply_table_formatting(ps_table)

    # Grand Total row in its own 1-row table
    gt_table = doc.add_table(rows=1, cols=3)
    gt_table.style = 'Table Grid'
    gt_row = gt_table.rows[0]
    gt_row.cells[0].paragraphs[0].clear()
    gt_row.cells[0].paragraphs[0].add_run('TOTAL')
    gt_row.cells[1].paragraphs[0].clear()
    gt_row.cells[1].paragraphs[0].add_run('')
    gt_row.cells[2].paragraphs[0].clear()
    gt_row.cells[2].paragraphs[0].add_run('{grandTotal}')
    # Bold the total row
    for cell in gt_row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    apply_table_formatting(gt_table, has_header=False)

    doc.add_page_break()

    # --- Section 5: Market Response (moved up — before Exec Summary) ---
    p_mr_open = doc.add_paragraph()
    p_mr_open.add_run('{#hasMarketResponses}')

    add_section_banner(doc, 'MARKET RESPONSE')

    mr_table = doc.add_table(rows=2, cols=4)
    mr_table.style = 'Table Grid'

    hdr = mr_table.rows[0]
    hdr.cells[0].text = 'Carrier'
    hdr.cells[1].text = 'Line of Business'
    hdr.cells[2].text = 'Status'
    hdr.cells[3].text = 'Notes / Reason'
    set_repeat_header(hdr)

    row = mr_table.rows[1]
    row.cells[0].paragraphs[0].clear()
    row.cells[0].paragraphs[0].add_run('{#marketResponses}{carrierName}')
    row.cells[1].paragraphs[0].clear()
    row.cells[1].paragraphs[0].add_run('{lobDisplay}')
    row.cells[2].paragraphs[0].clear()
    row.cells[2].paragraphs[0].add_run('{statusDisplay}')
    row.cells[3].paragraphs[0].clear()
    row.cells[3].paragraphs[0].add_run('{reason}{/marketResponses}')

    apply_table_formatting(mr_table)

    doc.add_page_break()

    p_mr_close = doc.add_paragraph()
    p_mr_close.add_run('{/hasMarketResponses}')

    # --- Section 6: Executive Summary ---
    add_section_banner(doc, 'EXECUTIVE SUMMARY')
    p_exec = doc.add_paragraph()
    p_exec.add_run('{narratives.executive_summary}')
    doc.add_page_break()

    # --- Section 6b: Special Notes (conditional) ---
    p_sn_open = doc.add_paragraph()
    p_sn_open.add_run('{#narratives.special_notes}')

    doc.add_heading('Special Notes', level=2)

    p_sn_content = doc.add_paragraph()
    p_sn_content.add_run('{narratives.special_notes}')

    p_sn_close = doc.add_paragraph()
    p_sn_close.add_run('{/narratives.special_notes}')

    # --- Section 7: Coverage Details (LOB injection point) ---
    add_section_banner(doc, 'COVERAGE DETAILS')
    p_lob = doc.add_paragraph()
    p_lob.add_run('{@lobSectionsXml}')
    doc.add_page_break()

    # --- Section 8: Recommendations ---
    add_section_banner(doc, 'RECOMMENDATIONS')
    p_rec = doc.add_paragraph()
    p_rec.add_run('{narratives.recommendations}')
    doc.add_page_break()

    # --- Section 9: Boilerplate injection ---
    p_bp = doc.add_paragraph()
    p_bp.add_run('{@boilerplateSectionsXml}')

    # --- Footer ---
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()

    pPr = footer_para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')

    tab1 = OxmlElement('w:tab')
    tab1.set(qn('w:val'), 'center')
    tab1.set(qn('w:pos'), '4680')
    tabs.append(tab1)

    tab2 = OxmlElement('w:tab')
    tab2.set(qn('w:val'), 'right')
    tab2.set(qn('w:pos'), '9360')
    tabs.append(tab2)
    pPr.append(tabs)

    run = footer_para.add_run('{amName}\t{proposalNumber}\t{generatedDate}')
    run.font.size = Pt(9)

    doc.save(output_path)
    inject_update_fields(output_path)
    print(f"  Created: {output_path}")


# ---------------------------------------------------------------------------
# LOB Partial: General Liability
# ---------------------------------------------------------------------------

def create_general_liability_docx(output_path):
    """Create general-liability.docx LOB partial template."""
    doc = Document()

    # 1. Section heading — dark blue banner
    add_section_banner(doc, '{sectionTitle}')

    # 2. Carrier info paragraphs
    add_carrier_info_paragraphs(doc)

    # Named Insured list
    add_named_insured_section(doc)

    # 3. Coverage Limits
    doc.add_heading('Coverage Limits', level=2)
    limits_table = doc.add_table(rows=2, cols=2)
    limits_table.style = 'Table Grid'
    limits_table.rows[0].cells[0].text = 'Coverage'
    limits_table.rows[0].cells[1].text = 'Limit'
    set_repeat_header(limits_table.rows[0])
    attr_row = limits_table.rows[1]
    attr_row.cells[0].paragraphs[0].clear()
    attr_row.cells[0].paragraphs[0].add_run('{#attributes}{label}')
    attr_row.cells[1].paragraphs[0].clear()
    attr_row.cells[1].paragraphs[0].add_run('{formattedValue}{/attributes}')
    apply_table_formatting(limits_table)

    # 4. Deductibles
    doc.add_heading('Deductibles', level=2)
    ded_table = doc.add_table(rows=2, cols=2)
    ded_table.style = 'Table Grid'
    ded_table.rows[0].cells[0].text = 'Deductible'
    ded_table.rows[0].cells[1].text = 'Details'
    set_repeat_header(ded_table.rows[0])
    ded_row = ded_table.rows[1]
    ded_row.cells[0].paragraphs[0].clear()
    ded_row.cells[0].paragraphs[0].add_run('{#deductibles}{deductibleType}')
    ded_row.cells[1].paragraphs[0].clear()
    ded_row.cells[1].paragraphs[0].add_run('{formattedValue}{/deductibles}')
    apply_table_formatting(ded_table)

    p_no_ded = doc.add_paragraph()
    p_no_ded.add_run('{^deductibles}None{/deductibles}')

    # 5. Additional Coverages
    doc.add_heading('Additional Coverages', level=2)
    cov_table = doc.add_table(rows=2, cols=3)
    cov_table.style = 'Table Grid'
    cov_table.rows[0].cells[0].text = 'Coverage'
    cov_table.rows[0].cells[1].text = 'Included'
    cov_table.rows[0].cells[2].text = 'Limit'
    set_repeat_header(cov_table.rows[0])
    cov_row = cov_table.rows[1]
    cov_row.cells[0].paragraphs[0].clear()
    cov_row.cells[0].paragraphs[0].add_run('{#coverages}{name}')
    cov_row.cells[1].paragraphs[0].clear()
    cov_row.cells[1].paragraphs[0].add_run('{#isIncluded}\u2713{/isIncluded}{^isIncluded}Excluded{/isIncluded}')
    cov_row.cells[2].paragraphs[0].clear()
    cov_row.cells[2].paragraphs[0].add_run('{limit}{/coverages}')
    apply_table_formatting(cov_table)

    # 6. Schedule of Locations — horizontal grid
    doc.add_heading('Schedule of Locations', level=2)
    loc_table = doc.add_table(rows=2, cols=4)
    loc_table.style = 'Table Grid'
    loc_headers = ['Location #', 'Street', 'City', 'State']
    for i, h in enumerate(loc_headers):
        loc_table.rows[0].cells[i].text = h
    set_repeat_header(loc_table.rows[0])
    loc_row = loc_table.rows[1]
    loc_keys = [
        '{#scheduleLocations}{itemNumber}',
        '{streetAddress}',
        '{city}',
        '{state}{/scheduleLocations}',
    ]
    for i, k in enumerate(loc_keys):
        loc_row.cells[i].paragraphs[0].clear()
        loc_row.cells[i].paragraphs[0].add_run(k)
    apply_table_formatting(loc_table)

    # 6b. Basis of Premium (GL classifications) — horizontal grid
    doc.add_heading('Basis of Premium', level=2)
    bop_table = doc.add_table(rows=2, cols=5)
    bop_table.style = 'Table Grid'
    bop_headers = ['Class Code', 'Description', 'Exposure', 'Rate', 'Premium']
    for i, h in enumerate(bop_headers):
        bop_table.rows[0].cells[i].text = h
    set_repeat_header(bop_table.rows[0])
    bop_row = bop_table.rows[1]
    bop_keys = [
        '{#glClassifications}{classCode}',
        '{classDescription}',
        '{exposure}',
        '{rate}',
        '{glPremium}{/glClassifications}',
    ]
    for i, k in enumerate(bop_keys):
        bop_row.cells[i].paragraphs[0].clear()
        bop_row.cells[i].paragraphs[0].add_run(k)
    apply_table_formatting(bop_table)

    # 7. Endorsements
    add_endorsements_table(doc)

    # 8. Premium
    p_premium = doc.add_paragraph()
    run = p_premium.add_run('Estimated Annual Premium: {premium}')
    run.bold = True

    # 9. Notes (conditional)
    add_notes_section(doc)

    doc.add_page_break()
    doc.save(output_path)
    print(f"  Created: {output_path}")


# ---------------------------------------------------------------------------
# LOB Partial: Workers' Compensation
# ---------------------------------------------------------------------------

def create_workers_comp_docx(output_path):
    """Create workers-compensation.docx LOB partial template."""
    doc = Document()

    # 1. Section heading — dark blue banner
    add_section_banner(doc, '{sectionTitle}')

    # 2. Carrier info paragraphs
    add_carrier_info_paragraphs(doc)

    # Named Insured list
    add_named_insured_section(doc)

    # 3. Coverage and Limits
    doc.add_heading('Coverage and Limits', level=2)
    limits_table = doc.add_table(rows=2, cols=2)
    limits_table.style = 'Table Grid'
    limits_table.rows[0].cells[0].text = 'Coverage'
    limits_table.rows[0].cells[1].text = 'Limit'
    set_repeat_header(limits_table.rows[0])
    attr_row = limits_table.rows[1]
    attr_row.cells[0].paragraphs[0].clear()
    attr_row.cells[0].paragraphs[0].add_run('{#attributes}{label}')
    attr_row.cells[1].paragraphs[0].clear()
    attr_row.cells[1].paragraphs[0].add_run('{formattedValue}{/attributes}')
    apply_table_formatting(limits_table)

    # 4. Employee Classification Schedule — horizontal grid
    doc.add_heading('Employee Classification Schedule', level=2)

    sched_table = doc.add_table(rows=2, cols=6)
    sched_table.style = 'Table Grid'
    wc_headers = ['State', 'Class Code', 'Description', 'Payroll', 'Rate', 'Premium']
    for i, h in enumerate(wc_headers):
        sched_table.rows[0].cells[i].text = h
    set_repeat_header(sched_table.rows[0])

    row = sched_table.rows[1]
    wc_keys = [
        '{#wcEmployeeClasses}{state}',
        '{classCode}',
        '{classDescription}',
        '{payroll}',
        '{ratePerHundred}',
        '{estimatedPremium}{/wcEmployeeClasses}',
    ]
    for i, k in enumerate(wc_keys):
        row.cells[i].paragraphs[0].clear()
        row.cells[i].paragraphs[0].add_run(k)

    apply_table_formatting(sched_table)

    # 5. Endorsements
    add_endorsements_table(doc)

    # 6. Premium
    p_premium = doc.add_paragraph()
    run = p_premium.add_run('Estimated Annual Premium: {premium} (subject to audit)')
    run.bold = True

    # 7. Notes (conditional)
    add_notes_section(doc)

    doc.add_page_break()
    doc.save(output_path)
    print(f"  Created: {output_path}")


# ---------------------------------------------------------------------------
# LOB Partial: Commercial Property
# ---------------------------------------------------------------------------

def create_commercial_property_docx(output_path):
    """Create commercial-property.docx LOB partial template."""
    doc = Document()

    # 1. Section heading — dark blue banner
    add_section_banner(doc, '{sectionTitle}')

    # 2. Carrier info paragraphs (Property adds policy # field)
    add_carrier_info_paragraphs(doc, fields=[
        ('CARRIER', '{carrier.name}'),
        ('A.M. BEST RATING', '{carrier.amBestRating}'),
        ('POLICY PERIOD', '{effectiveDate} \u2014 {expirationDate}'),
        ('COVERAGE', '{sectionTitle}'),
        ('POLICY #', '{#policyNumber}{policyNumber}{/policyNumber}{^policyNumber}\u2014{/policyNumber}'),
    ])

    # Named Insured list
    add_named_insured_section(doc)

    # 3. Property Details
    doc.add_heading('Property Details', level=2)
    limits_table = doc.add_table(rows=2, cols=2)
    limits_table.style = 'Table Grid'
    limits_table.rows[0].cells[0].text = 'Coverage'
    limits_table.rows[0].cells[1].text = 'Limit'
    set_repeat_header(limits_table.rows[0])
    attr_row = limits_table.rows[1]
    attr_row.cells[0].paragraphs[0].clear()
    attr_row.cells[0].paragraphs[0].add_run('{#attributes}{label}')
    attr_row.cells[1].paragraphs[0].clear()
    attr_row.cells[1].paragraphs[0].add_run('{formattedValue}{/attributes}')
    apply_table_formatting(limits_table)

    # 4. Deductibles
    doc.add_heading('Deductibles', level=2)
    ded_table = doc.add_table(rows=2, cols=2)
    ded_table.style = 'Table Grid'
    ded_table.rows[0].cells[0].text = 'Deductible Type'
    ded_table.rows[0].cells[1].text = 'Details'
    set_repeat_header(ded_table.rows[0])
    ded_row = ded_table.rows[1]
    ded_row.cells[0].paragraphs[0].clear()
    ded_row.cells[0].paragraphs[0].add_run('{#deductibles}{deductibleType}')
    ded_row.cells[1].paragraphs[0].clear()
    ded_row.cells[1].paragraphs[0].add_run('{formattedValue}{/deductibles}')
    apply_table_formatting(ded_table)

    # 5. Location / Building Schedule — horizontal 5-col grid
    doc.add_heading('Location / Building Schedule', level=2)
    prop_sched_table = doc.add_table(rows=2, cols=5)
    prop_sched_table.style = 'Table Grid'
    prop_sched_headers = ['Location #', 'Building #', 'Description', 'Address', 'Building Limit']
    for i, h in enumerate(prop_sched_headers):
        prop_sched_table.rows[0].cells[i].text = h
    set_repeat_header(prop_sched_table.rows[0])
    prop_row = prop_sched_table.rows[1]
    prop_keys = [
        '{#propertySchedule}{locationNumber}',
        '{buildingNumber}',
        '{description}',
        '{address}',
        '{buildingLimit}{/propertySchedule}',
    ]
    for i, k in enumerate(prop_keys):
        prop_row.cells[i].paragraphs[0].clear()
        prop_row.cells[i].paragraphs[0].add_run(k)
    apply_table_formatting(prop_sched_table)

    # 6. Endorsements
    add_endorsements_table(doc)

    # 7. Premium
    p_premium = doc.add_paragraph()
    run = p_premium.add_run('Estimated Annual Premium: {premium} | Total Cost: {totalCost}')
    run.bold = True

    # 8. Notes (conditional)
    add_notes_section(doc)

    doc.add_page_break()
    doc.save(output_path)
    print(f"  Created: {output_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    base_dir = Path(__file__).resolve().parent.parent
    templates_dir = base_dir / 'templates'

    # Create output directories
    (templates_dir / 'verticals' / 'nba').mkdir(parents=True, exist_ok=True)
    (templates_dir / 'lob-partials').mkdir(parents=True, exist_ok=True)
    (templates_dir / 'registry').mkdir(parents=True, exist_ok=True)
    (base_dir / 'docs').mkdir(parents=True, exist_ok=True)

    print("Generating proposal templates...")

    # Generate .docx files
    create_master_docx(str(templates_dir / 'verticals' / 'nba' / 'master.docx'))
    create_general_liability_docx(str(templates_dir / 'lob-partials' / 'general-liability.docx'))
    create_workers_comp_docx(str(templates_dir / 'lob-partials' / 'workers-compensation.docx'))
    create_commercial_property_docx(str(templates_dir / 'lob-partials' / 'commercial-property.docx'))

    # meta.json
    meta_path = templates_dir / 'verticals' / 'nba' / 'meta.json'
    meta = {
        "templateId": "nba-v1",
        "vertical": "nba",
        "displayName": "NBA \u2014 Standard Proposal",
        "version": "1.0.0",
        "s3Key": "verticals/nba/master.docx",
        "lobPartials": ["GeneralLiability", "WorkersCompensation", "CommercialProperty"],
        "defaultBoilerplate": ["about_fortress", "am_best_disclaimer", "e_o_disclosure"],
        "active": True
    }
    meta_path.write_text(json.dumps(meta, indent=2) + '\n')
    print(f"  Created: {meta_path}")

    # boilerplate.json
    bp_path = templates_dir / 'registry' / 'boilerplate.json'
    boilerplate = {
        "version": "1.0.0",
        "blocks": {
            "about_fortress": {
                "id": "about_fortress",
                "displayName": "About NBA Insurance Services",
                "type": "wordml",
                "content": '<w:p><w:pPr><w:pStyle w:val="Heading1"/></w:pPr><w:r><w:t>About NBA Insurance Services</w:t></w:r></w:p><w:p><w:r><w:t>NBA Insurance Services is a full-service independent insurance agency serving commercial clients across the Southeast. Our team of experienced professionals delivers tailored risk management solutions and dedicated service. We are proud to serve {insuredName} and look forward to a continued partnership.</w:t></w:r></w:p>',
                "variables": ["insuredName"],
                "notes": "Standard About Fortress section."
            },
            "am_best_disclaimer": {
                "id": "am_best_disclaimer",
                "displayName": "AM Best Rating Disclaimer",
                "type": "text",
                "content": "Carrier financial strength ratings are provided by A.M. Best Company as of the date of this proposal and are subject to change. NBA Insurance Services makes no representations regarding carrier solvency.",
                "variables": [],
                "notes": "Required on all proposals."
            },
            "e_o_disclosure": {
                "id": "e_o_disclosure",
                "displayName": "E&O Disclosure",
                "type": "text",
                "content": "This proposal is prepared for the exclusive use of {insuredName} and is based on information provided by the insured. Coverage is not bound until confirmed in writing by the issuing carrier.",
                "variables": ["insuredName"],
                "notes": "Required on all proposals."
            }
        }
    }
    bp_path.write_text(json.dumps(boilerplate, indent=2) + '\n')
    print(f"  Created: {bp_path}")

    # upload-to-s3.sh
    upload_path = templates_dir / 'upload-to-s3.sh'
    upload_script = '''#!/bin/bash
# Upload all proposal-generator templates to S3
BUCKET="fortress-tools"
PREFIX="fip-proposal-templates"
TEMPLATES_DIR="$(dirname "$0")"

aws s3 cp "$TEMPLATES_DIR/verticals/nba/master.docx" "s3://$BUCKET/$PREFIX/verticals/nba/master.docx" --profile fortress-tools-deployer
aws s3 cp "$TEMPLATES_DIR/verticals/nba/meta.json"   "s3://$BUCKET/$PREFIX/verticals/nba/meta.json"   --profile fortress-tools-deployer
aws s3 cp "$TEMPLATES_DIR/lob-partials/general-liability.docx"    "s3://$BUCKET/$PREFIX/lob-partials/general-liability.docx"    --profile fortress-tools-deployer
aws s3 cp "$TEMPLATES_DIR/lob-partials/workers-compensation.docx" "s3://$BUCKET/$PREFIX/lob-partials/workers-compensation.docx" --profile fortress-tools-deployer
aws s3 cp "$TEMPLATES_DIR/lob-partials/commercial-property.docx"  "s3://$BUCKET/$PREFIX/lob-partials/commercial-property.docx"  --profile fortress-tools-deployer
aws s3 cp "$TEMPLATES_DIR/registry/boilerplate.json" "s3://$BUCKET/$PREFIX/registry/boilerplate.json" --profile fortress-tools-deployer
echo "Done."
'''
    upload_path.write_text(upload_script)
    upload_path.chmod(0o755)
    print(f"  Created: {upload_path}")

    # docs/template-authoring-checklist.md
    checklist_path = base_dir / 'docs' / 'template-authoring-checklist.md'
    checklist = '''# Template Authoring Checklist

## Prerequisites

- Python 3.8+ with `python-docx` installed (`pip install python-docx`)
- Run from: `/home/fredw/projects/fip/services/proposal-generator`
- Command: `python3 scripts/generate-templates.py`

## Single-Run Rule

Every `{tag}` must be contained in ONE Word run. Never split a tag across formatting boundaries.

**How to verify:** Open the `.docx` as a ZIP, inspect `word/document.xml`, and confirm each `{...}` tag appears within a single `<w:t>` element. python-docx programmatic creation guarantees this when you set the full tag string in one `add_run()` call.

**Do NOT:**
```python
p.add_run('{insured')
p.add_run('Name}')  # BROKEN: tag split across two runs
```

**Do:**
```python
p.add_run('{insuredName}')  # Correct: single run
```

## Table Formatting Requirements

Every table must have these Word XML properties applied:

| Property | Where | Purpose |
|----------|-------|---------|
| `w:keepNext` | All paragraph runs in rows (except last row) | Prevents orphan rows across page breaks |
| `w:tblHeader` | Header rows | Repeats header on each page |
| `w:keepLines` | All table cells | Keeps cell content together |

Use `apply_table_formatting(table, has_header=True)` to apply all at once.

## Loop Syntax Patterns

### Table Row Loops

Place `{#array}` in the **first cell** and `{/array}` in the **last cell** of the row to repeat.

```
| {#items}{name} | {value}{/items} |   <- entire row repeats
```

### Paragraph Loops (paragraphLoop: true)

Place `{#array}` and `{/array}` on their **own standalone paragraphs**. Content paragraphs go between them.

```
{#scheduleItems}           <- own paragraph
Item {itemNumber}          <- content paragraph
{/scheduleItems}           <- own paragraph
```

### Nested Loops

Outer loop uses paragraph style, inner loop can use table row style:

```
{#premiumSummary.byLob}                    <- paragraph
  | {#quotes}{carrier} | {premium}{/quotes} |  <- table row
{/premiumSummary.byLob}                    <- paragraph
```

## Conditional Section Patterns

- Truthy: `{#field}...{/field}` — renders when field is truthy
- Falsy: `{^field}...{/field}` — renders when field is null/false/empty
- Inline conditional: `{#isAdmitted}Admitted{/isAdmitted}{^isAdmitted}Surplus Lines{/isAdmitted}`

## Raw XML Injection

`{@lobSectionsXml}` and `{@boilerplateSectionsXml}` must each be in their **own dedicated paragraph** with **NO other text**.

```python
p = doc.add_paragraph()
p.add_run('{@lobSectionsXml}')  # Nothing else in this paragraph
```

## Data Contract Notes

- `carrier` is an **OBJECT** — use `{carrier.name}`, `{carrier.amBestRating}`
- `deductibles[].formattedValue` — use `{formattedValue}` for display (shows "$25,000" for flat, "5%" for percentage)
- `scheduleItems[].children` may be null (e.g., WC) — only include children loop if LOB has nested items
- `premiumSummary.byLob[].quotes[].carrier` is a string (carrier name), not an object

## Verification

After generating templates:

1. Check file sizes are > 0: `ls -la templates/**/*.docx`
2. Validate ZIP structure: `python3 -c "import zipfile; zipfile.ZipFile(\'templates/verticals/nba/master.docx\').testzip()"`
3. Inspect XML for split runs: unzip and search for split `{` / `}` across `<w:t>` elements
'''
    checklist_path.write_text(checklist)
    print(f"  Created: {checklist_path}")

    # Write .generated marker
    gen_path = templates_dir / '.generated'
    now = datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')
    generated_files = [
        'templates/verticals/nba/master.docx',
        'templates/verticals/nba/meta.json',
        'templates/lob-partials/general-liability.docx',
        'templates/lob-partials/workers-compensation.docx',
        'templates/lob-partials/commercial-property.docx',
        'templates/registry/boilerplate.json',
        'templates/upload-to-s3.sh',
        'docs/template-authoring-checklist.md',
    ]
    gen_path.write_text(f"Generated: {now}\n\nFiles:\n" + '\n'.join(f"  - {f}" for f in generated_files) + '\n')
    print(f"  Created: {gen_path}")

    print("\nSUCCESS: All templates generated")


if __name__ == '__main__':
    main()
