#!/usr/bin/env python3
"""
Patch master.docx with updated tags, labels, branding, and content changes.
Output: docs/master-updated.docx
"""

import copy
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'


# ---------------------------------------------------------------------------
# Text replacement helpers
# ---------------------------------------------------------------------------

def replace_in_para(para, old, new):
    """Replace text in a paragraph. Tries per-run first to preserve formatting,
    falls back to collapsing runs (for docxtemplater tags split across runs)."""
    found = False
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            found = True
    if found:
        return True
    full = ''.join(r.text for r in para.runs)
    if old in full:
        new_full = full.replace(old, new)
        if para.runs:
            para.runs[0].text = new_full
            for r in para.runs[1:]:
                r.text = ''
        return True
    return False


def replace_in_cell(cell, old, new):
    for para in cell.paragraphs:
        replace_in_para(para, old, new)


def replace_in_table(table, old, new):
    for row in table.rows:
        for cell in row.cells:
            replace_in_cell(cell, old, new)


def replace_in_header_footer(section, old, new):
    for part in [
        section.header, section.even_page_header, section.first_page_header,
        section.footer, section.even_page_footer, section.first_page_footer,
    ]:
        if part is None:
            continue
        try:
            linked = part.is_linked_to_previous
        except Exception:
            linked = False
        if linked:
            continue
        for para in part.paragraphs:
            replace_in_para(para, old, new)
        for tbl in part.tables:
            replace_in_table(tbl, old, new)


def replace_everywhere(doc, old, new):
    """Replace old→new in body paragraphs, all tables, and all headers/footers."""
    for para in doc.paragraphs:
        replace_in_para(para, old, new)
    for table in doc.tables:
        replace_in_table(table, old, new)
    for section in doc.sections:
        replace_in_header_footer(section, old, new)


# ---------------------------------------------------------------------------
# XML helper: make a simple run element
# ---------------------------------------------------------------------------

def make_run(text, bold=False, italic=False, color=None, sz=None, font=None):
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if font:
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font)
        rFonts.set(qn('w:hAnsi'), font)
        rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement('w:b'))
    if italic:
        i_el = OxmlElement('w:i')
        rPr.append(i_el)
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    if sz:
        s = OxmlElement('w:sz')
        s.set(qn('w:val'), str(sz))
        rPr.append(s)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        t.set(XML_SPACE, 'preserve')
    r.append(t)
    return r


def make_paragraph(runs_spec, pPr_source=None):
    """Create a <w:p> element. runs_spec: list of (text, bold, italic) tuples."""
    p = OxmlElement('w:p')
    if pPr_source is not None:
        pPr = pPr_source.find(f'{{{NS}}}pPr')
        if pPr is not None:
            p.append(copy.deepcopy(pPr))
    for spec in runs_spec:
        text, bold, italic = spec
        p.append(make_run(text, bold=bold, italic=italic))
    return p


# ---------------------------------------------------------------------------
# Load document
# ---------------------------------------------------------------------------

doc = docx.Document('docs/master.docx')

# ---------------------------------------------------------------------------
# A. Premium Summary table (Table 2) changes
# ---------------------------------------------------------------------------

table2 = doc.tables[2]

# A1. Rename "Base Premium" → "Estimated Manual Premium"
#     Update tag {basePremium} → {manualPremium}
row1 = table2.rows[1]
for cell in row1.cells:
    for para in cell.paragraphs:
        replace_in_para(para, 'Base Premium', 'Estimated Manual Premium')
        replace_in_para(para, '{basePremium}', '{manualPremium}')

# A2. Insert EMR conditional row after row 1
tr_base = table2.rows[1]._tr
emr_tr = copy.deepcopy(tr_base)
cells_in_emr = emr_tr.findall(f'{{{NS}}}tc')

if len(cells_in_emr) >= 2:
    # Cell 0: label with opening conditional tag
    for p in cells_in_emr[0].findall(f'.//{{{NS}}}p'):
        for t_el in p.findall(f'.//{{{NS}}}t'):
            t_el.text = '{#hasEmr}Experience Modification Rate'
    # Cell 1: value + closing conditional tag
    for p in cells_in_emr[1].findall(f'.//{{{NS}}}p'):
        for t_el in p.findall(f'.//{{{NS}}}t'):
            t_el.text = '{emr}{/hasEmr}'

# Insert EMR row immediately after row 1
tr_base.addnext(emr_tr)

# A3. Update {estPremium} → {annualPremium}  (Estimated Annual Premium value cell)
# Table reference may have shifted; iterate all rows
for row in doc.tables[2].rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            replace_in_para(para, '{estPremium}', '{annualPremium}')

# A4. Update {totalEstimatedPremium} → {totalEstimatedCost} in Table 2 ONLY
for row in doc.tables[2].rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            replace_in_para(para, '{totalEstimatedPremium}', '{totalEstimatedCost}')

# ---------------------------------------------------------------------------
# B. Class Schedule table (Table 7) — total row
# ---------------------------------------------------------------------------

table7 = doc.tables[7]
for row in table7.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            # Rename label
            replace_in_para(para, 'Total Estimated Premium', 'Total Estimated Manual Premium')
            # Update tag to {manualPremium} (was {totalEstimatedPremium})
            replace_in_para(para, '{totalEstimatedPremium}', '{manualPremium}')

# ---------------------------------------------------------------------------
# C. Move Surplus Contribution block to page 2, above "What's next"
# ---------------------------------------------------------------------------

# Locate paragraphs by content
whatsnext_para = None
surplus_heading_para = None
surplus_body_para = None

for p in doc.paragraphs:
    txt = p.text.strip()
    if txt in ("What’s next", "What's next"):
        whatsnext_para = p
    elif txt == 'Surplus Contribution':
        surplus_heading_para = p
    elif txt.startswith('As a self-insured group (SIG), BAWNSIG'):
        surplus_body_para = p

if whatsnext_para and surplus_heading_para and surplus_body_para:
    # Deep-copy the XML elements for insertion
    heading_copy = copy.deepcopy(surplus_heading_para._p)
    body_copy = copy.deepcopy(surplus_body_para._p)

    # Insert before "What's next": heading first, then body after heading
    whatsnext_para._p.addprevious(heading_copy)
    heading_copy.addnext(body_copy)

    # Remove originals from their current location
    surplus_heading_para._p.getparent().remove(surplus_heading_para._p)
    surplus_body_para._p.getparent().remove(surplus_body_para._p)
    print("C: Surplus Contribution block moved to page 2")
else:
    print("WARNING C: Could not locate surplus block or What's next paragraph")
    print(f"  whatsnext_para: {whatsnext_para}")
    print(f"  surplus_heading_para: {surplus_heading_para}")
    print(f"  surplus_body_para: {surplus_body_para}")

# ---------------------------------------------------------------------------
# D. Capitalize "Program highlights" → "Program Highlights"
# ---------------------------------------------------------------------------

for para in doc.paragraphs:
    if para.text.strip() == 'Program highlights':
        for run in para.runs:
            run.text = run.text.replace('Program highlights', 'Program Highlights')
        print("D: Program highlights capitalized")
        break

# ---------------------------------------------------------------------------
# E. Contact / address block
# ---------------------------------------------------------------------------

replace_everywhere(doc, '(775) 555-0100', '877-888-4140')
replace_everywhere(doc, 'dslater@nbais.com', 'dslater@nbainsurancesolutions.com')
replace_everywhere(doc, 'www.nbais.com', 'nbainsurancesolutions.com')
replace_everywhere(doc, "1234 Builder’s Way, Reno, NV 89501", '603 South Carson Street, Carson City, NV 89701')
replace_everywhere(doc, "1234 Builder's Way, Reno, NV 89501", '603 South Carson Street, Carson City, NV 89701')
print("E: Contact info updated")

# ---------------------------------------------------------------------------
# F. Branding — Services → Solutions + LLC
# ---------------------------------------------------------------------------

# F13. Replace "Services" → "Solutions" throughout
replace_everywhere(doc, 'Services', 'Solutions')
print("F13: Services → Solutions applied")

# F14. Cover page (P0): add ", LLC" after "Nevada Builders Alliance Insurance Solutions"
#      Only the first occurrence (cover page heading)
for para in doc.paragraphs:
    txt = para.text
    if 'Nevada Builders Alliance Insurance Solutions' in txt and 'LLC' not in txt:
        for run in para.runs:
            if 'Nevada Builders Alliance Insurance Solutions' in run.text:
                run.text = run.text.replace(
                    'Nevada Builders Alliance Insurance Solutions',
                    'Nevada Builders Alliance Insurance Solutions, LLC'
                )
        print(f"F14: LLC added to cover page paragraph: {para.text[:80]}")
        break

# F15. Page 6 disclaimer: "NBAIS is an insurance program" → "NBA Insurance Solutions, LLC is..."
for para in doc.paragraphs:
    if para.text.startswith('This proposal is not a binder'):
        for run in para.runs:
            if 'NBAIS is an insurance program' in run.text:
                run.text = run.text.replace(
                    'NBAIS is an insurance program',
                    'NBA Insurance Solutions, LLC is an insurance program'
                )
        print(f"F15: LLC added to disclaimer: {para.text[:80]}")
        break

# ---------------------------------------------------------------------------
# G. Excluded persons Note replacement (P35)
# ---------------------------------------------------------------------------

for para in doc.paragraphs:
    if para.text.startswith('Note: Officers or members electing to reject'):
        p_el = para._p

        # Remove all existing runs
        for r in list(p_el.findall(f'{{{NS}}}r')):
            p_el.remove(r)

        # Build new content for this paragraph:
        # Bold+italic "D-43 form required: " + italic body
        d43_lead = 'D-43 form required: '
        d43_body = ('A completed and signed Form D-43 is required for each individual listed '
                    'above prior to binding. Coverage cannot be excluded without a signed '
                    'D-43 on file.')
        p_el.append(make_run(d43_lead, bold=True, italic=True))
        p_el.append(make_run(d43_body, bold=False, italic=True))

        # Insert a new paragraph immediately after for "Important:" section
        imp_lead = 'Important: '
        imp_body = ('Excluded individuals are not covered under this workers’ '
                    'compensation policy. Please confirm these elections are accurate '
                    'prior to binding.')

        imp_p = OxmlElement('w:p')
        pPr = p_el.find(f'{{{NS}}}pPr')
        if pPr is not None:
            imp_p.append(copy.deepcopy(pPr))
        imp_p.append(make_run(imp_lead, bold=True, italic=True))
        imp_p.append(make_run(imp_body, bold=False, italic=True))

        p_el.addnext(imp_p)
        print("G: Note replaced with D-43 form required text")
        break

# ---------------------------------------------------------------------------
# Save output
# ---------------------------------------------------------------------------

doc.save('docs/master-updated.docx')
print("\nSaved: docs/master-updated.docx")
